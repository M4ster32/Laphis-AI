"""Chapter 6 — Método e Planeamento."""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...docx_utils.blocks import add_heading, add_paragraph, add_bullet
from ...docx_utils.theme import GREY


def render(doc):
    """Render chapter 6 with methodology notes and milestone list."""
    add_heading(doc, "Método e Planeamento", level=1, number="6")

    add_paragraph(doc,
        "A metodologia adotada é híbrida, combinando elementos de "
        "metodologias ágeis (desenvolvimento iterativo, entregas "
        "frequentes de incrementos funcionais) com fases clássicas em "
        "cascata para a documentação — levantamento de requisitos, "
        "modelação e testes finais. O trabalho individual exclui a "
        "necessidade de cerimónias formais de equipa, mas as práticas "
        "de controlo de versões (branches, commits atómicos, mensagens "
        "convencionais) foram mantidas com rigor.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "O desenvolvimento foi dividido em dois semestres letivos, "
        "correspondentes às UCs Projeto I (1º semestre, 2025/2026) e "
        "Projeto II (2º semestre, 2025/2026). Os marcos principais são:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

    add_paragraph(doc, "Projeto I — 1º Semestre (Set 2025 – Jan 2026)",
                  bold=True, space_after=4)
    add_bullet(doc, "Set–Out 2025: Levantamento do estado da arte, estudo "
                    "de soluções concorrentes, definição do problema.")
    add_bullet(doc, "Nov 2025: Entrega intercalar — especificação de "
                    "requisitos e modelação preliminar.")
    add_bullet(doc, "Dez 2025 – Jan 2026: Implementação do núcleo "
                    "funcional (autenticação, tracking básico, primeiros "
                    "endpoints).")
    add_bullet(doc, "Jan 2026: Entrega final de Projeto I.")

    add_paragraph(doc, "Projeto II — 2º Semestre (Fev 2026 – Mai 2026)",
                  bold=True, space_after=4)
    add_bullet(doc, "Fev 2026: Expansão do motor de recomendação, "
                    "integração com OpenAI e pipeline RAG inicial.")
    add_bullet(doc, "Mar 2026: Desenvolvimento do frontend completo, "
                    "sistema de design premium, dark/light mode.")
    add_bullet(doc, "Abr 2026: Testes automatizados, modernização do "
                    "código (remoção de deprecations), deploy em "
                    "produção. Entrega intercalar (o presente relatório).")
    add_bullet(doc, "Mai 2026: Validação por terceiros, refinamentos "
                    "finais, documentação completa. Entrega final e "
                    "defesa.")

    add_paragraph(doc,
        "A Figura 7 apresenta o cronograma Gantt detalhado. À data desta "
        "entrega intercalar, todos os marcos até ao final de abril foram "
        "cumpridos dentro do prazo planeado. As dificuldades mais "
        "relevantes encontradas foram:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    for item in [
        "Migração do SQLAlchemy v1 para v2 — implicou refactor da camada "
        "de dados, mas produziu código mais limpo e tipado.",
        "Modernização para Pydantic v2 — alteração de @validator para "
        "@model_validator e de orm_mode para ConfigDict.",
        "Dimensionamento do modelo de dados — iteração sobre o diagrama "
        "ER com base em casos de uso concretos.",
        "Configuração CORS em produção — ajuste de origens permitidas "
        "para o domínio Vercel.",
    ]:
        add_bullet(doc, item)

    add_paragraph(doc,
        "[Espaço reservado para Figura 7 — Cronograma Gantt. A gerar em "
        "ferramenta externa (TeamGantt, MS Project) e incluir antes da "
        "entrega final.]",
        italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

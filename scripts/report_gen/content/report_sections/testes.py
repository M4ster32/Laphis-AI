"""Chapter 5 — Testes e Validação."""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...docx_utils.blocks import (
    add_heading,
    add_paragraph,
    add_bullet,
    add_caption,
)
from ...docx_utils.tables import add_table
from ..project_facts import TEST_MODULES, METRICS


def render(doc):
    """Render chapter 5 with the test plan table."""
    add_heading(doc, "Testes e Validação", level=1, number="5")

    add_paragraph(doc,
        "O plano de testes do LAPHIS é construído em dois níveis "
        "complementares: testes automatizados do backend, focados em "
        "correção funcional e regressão, e testes de aceitação por "
        "utilizadores, orientados à validação operacional em contexto "
        "real. O plano de testes foi desenhado em paralelo com a "
        "especificação dos requisitos (Secção 3.1), de modo a garantir "
        "cobertura dos comportamentos críticos do sistema.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc, "Testes automatizados — cobertura atual",
                  bold=True, space_after=4)
    add_paragraph(doc,
        f"A suite atual reúne {METRICS['test_count']} testes "
        "automatizados, executados via pytest com cliente HTTP "
        "assíncrono httpx. Cada teste corre contra uma base de dados "
        "SQLite in-memory criada e destruída por teste, garantindo "
        "isolamento total.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_caption(doc, "Plano de testes por módulo.", kind="Tabela 7")
    add_table(doc,
        ["Ficheiro", "Testes", "Cobertura funcional"],
        [list(row) for row in TEST_MODULES],
        col_widths=[3.5, 2.0, 10.5])

    add_paragraph(doc, "Testes de aceitação por terceiros (planeados)",
                  bold=True, space_after=4)
    add_paragraph(doc,
        "Para a entrega final será aplicado um inquérito de satisfação a "
        "10-15 utilizadores-teste, combinando respostas Likert e "
        "perguntas abertas, com o guião detalhado no Anexo de testes. "
        "Os critérios de avaliação incluem:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    for item in [
        "Usabilidade — facilidade de registo, navegação, realização das "
        "tarefas principais.",
        "Percepção de utilidade — os utilizadores sentem que o LAPHIS "
        "resolve o problema de fragmentação identificado?",
        "Qualidade das recomendações — os planos e sugestões do "
        "assistente IA são percebidos como personalizados?",
        "Desempenho percebido — tempo de resposta aceitável?",
        "Disponibilidade de funcionalidades — falta alguma funcionalidade "
        "considerada essencial?",
    ]:
        add_bullet(doc, item)

    add_paragraph(doc,
        "A análise de risco é realizada com recurso a diagrama de "
        "causa-efeito (Ishikawa), identificando os principais modos de "
        "falha: (i) indisponibilidade do backend (mitigação: health "
        "checks e auto-restart), (ii) expiração de token JWT em operação "
        "longa (mitigação: refresh automático no cliente), (iii) falha na "
        "API OpenAI (mitigação: fallback para modo regras), (iv) limites "
        "de tier gratuito da base de dados (mitigação: monitorização de "
        "tamanho e limpeza de sessões expiradas).",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

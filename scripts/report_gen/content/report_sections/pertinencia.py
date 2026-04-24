"""Chapter 2 — Pertinência e Viabilidade."""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...docx_utils.blocks import (
    add_heading,
    add_paragraph,
    add_bullet,
    add_caption,
)
from ...docx_utils.tables import add_table


def render(doc):
    """Render chapter 2 with benchmarking table."""
    add_heading(doc, "Pertinência e Viabilidade", level=1, number="2")

    add_heading(doc, "Pertinência", level=2, number="2.1")
    add_paragraph(doc,
        "O impacto positivo esperado do LAPHIS assenta em três pilares. "
        "Primeiro, redução do custo cognitivo de gestão da saúde pessoal: "
        "ao consolidar múltiplos domínios numa única aplicação, "
        "elimina-se a necessidade de alternar entre ferramentas e "
        "reconciliar manualmente dados dispersos. Segundo, personalização "
        "real: um motor de recomendação que cruza perfil biométrico, "
        "objetivos declarados e registos históricos gera aconselhamento "
        "adaptado, em contraste com os planos genéricos oferecidos pela "
        "maioria das soluções de mercado. Terceiro, acessibilidade "
        "tecnológica: o sistema é entregue como aplicação web, "
        "dispensando instalação e funcionando em qualquer dispositivo com "
        "browser moderno.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "A validação por terceiros desta pertinência foi iniciada de "
        "forma informal junto de colegas e utilizadores-teste durante o "
        "desenvolvimento. Nesta fase intercalar, o feedback recolhido "
        "será sistematizado e apresentado em inquérito estruturado na "
        "entrega final.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Viabilidade", level=2, number="2.2")
    add_paragraph(doc,
        "A viabilidade técnica do projeto está demonstrada pelo estado "
        "atual de implementação: a aplicação encontra-se funcional em "
        "ambiente de desenvolvimento, com todos os módulos principais "
        "(autenticação, tracking, planos, chat) operacionais, e o deploy "
        "em produção em execução à data desta entrega intercalar. A "
        "utilização de ferramentas gratuitas ou com tiers gratuitos "
        "generosos (Vercel, Render, GitHub, PostgreSQL free tier) garante "
        "a sustentabilidade financeira do projeto enquanto trabalho "
        "académico. A componente de IA é opcional: na ausência de chave "
        "OpenAI, o sistema funciona integralmente em modo baseado em "
        "regras.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "Do ponto de vista económico, um eventual cenário comercial "
        "apoia-se em modelo freemium — funcionalidades essenciais "
        "gratuitas, funcionalidades premium (assistente IA avançado, "
        "integração com wearables, histórico ilimitado) mediante "
        "subscrição. A manutenção e evolução após o término do TFC é "
        "viável, dado o caráter modular da arquitetura e a cobertura de "
        "testes existente.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Análise Comparativa com Soluções Existentes",
                level=2, number="2.3")

    add_heading(doc, "Soluções existentes", level=3, number="2.3.1")
    add_paragraph(doc,
        "Foram identificadas no mercado quatro soluções principais, cada "
        "uma especializada num subconjunto do domínio de saúde e fitness:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    add_bullet(doc,
        "MyFitnessPal — tracking de nutrição com base de dados alimentar "
        "extensa; funcionalidades de treino e integração limitadas.")
    add_bullet(doc,
        "Strava — rede social e tracking de atividade física outdoor; sem "
        "componente de nutrição ou bem-estar mental.")
    add_bullet(doc,
        "Calm / Headspace — meditação guiada e mindfulness; sem "
        "integração com tracking físico ou nutricional.")
    add_bullet(doc,
        "Fitbit / Apple Health — hubs de dados biométricos agregados por "
        "dispositivos vestíveis; fracos em recomendação e planeamento.")

    add_heading(doc, "Análise de benchmarking", level=3, number="2.3.2")
    add_paragraph(doc,
        "A Tabela 1 sistematiza a comparação da solução LAPHIS com as "
        "alternativas identificadas, sinalizando com \"X\" a presença de "
        "cada característica-chave.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_caption(doc,
        "Análise de benchmarking comparativo com soluções existentes.",
        kind="Tabela 1")
    add_table(doc,
        ["Característica", "MyFitnessPal", "Strava", "Calm", "Fitbit",
         "LAPHIS"],
        [
            ["Tracking treino", "", "X", "", "X", "X"],
            ["Tracking nutrição", "X", "", "", "parcial", "X"],
            ["Hidratação", "parcial", "", "", "parcial", "X"],
            ["Peso corporal", "X", "", "", "X", "X"],
            ["Meditação", "", "", "X", "parcial", "X"],
            ["Planos personalizados", "parcial", "parcial", "", "", "X"],
            ["Chat IA", "", "", "", "", "X"],
            ["RAG sobre PDFs", "", "", "", "", "X"],
            ["Dark mode nativo", "parcial", "X", "X", "X", "X"],
            ["Open source / académico", "", "", "", "", "X"],
        ],
        col_widths=[3.5, 2.5, 2.0, 2.0, 2.0, 2.0])

    add_heading(doc, "Proposta de Inovação e Mais-Valias", level=2,
                number="2.4")
    add_paragraph(doc,
        "O elemento distintivo do LAPHIS é a integração cross-domínio "
        "combinada com uma camada de IA opcional e um motor de regras "
        "transparente. A grande maioria das soluções comerciais opera num "
        "único silo, deixando ao utilizador a responsabilidade de cruzar "
        "mentalmente os dados. O LAPHIS inverte este paradigma: os "
        "registos de treino influenciam as recomendações nutricionais, o "
        "peso corporal alimenta o cálculo de hidratação recomendada, e o "
        "feedback sobre planos informa ajustes futuros.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "Mais-valias adicionais: (i) transparência do motor de "
        "recomendação — as fórmulas usadas (Harris-Benedict, distribuição "
        "de macros) são publicadas no código, ao contrário do caráter "
        "opaco da maioria das aplicações; (ii) pipeline RAG sobre PDFs, "
        "que permite ao utilizador carregar documentos próprios "
        "(relatórios médicos, planos nutricionais de profissionais) e "
        "fazer consultas em linguagem natural; (iii) design system "
        "premium com dark/light mode completo; (iv) disponibilização em "
        "ambiente produtivo real.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Identificação de Oportunidade de Negócio", level=2,
                number="2.5")
    add_paragraph(doc,
        "Caso se pretenda transpor o LAPHIS para exploração comercial, o "
        "modelo de negócio natural é freemium com segmento B2C e extensão "
        "B2B a prazo:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    add_bullet(doc, "Tier gratuito — tracking completo, chat com IA "
                    "limitado a N sessões/mês, planos gerados por regras.")
    add_bullet(doc, "Tier Premium (€4,99/mês) — chat IA ilimitado, planos "
                    "gerados por LLM, integração com wearables, histórico "
                    "ilimitado.")
    add_bullet(doc, "Tier Profissional — marketplace de planos validados "
                    "por nutricionistas e personal trainers certificados, "
                    "com comissão sobre transações.")
    add_bullet(doc, "Licenciamento B2B — ginásios e clínicas podem "
                    "oferecer LAPHIS white-label aos seus clientes.")

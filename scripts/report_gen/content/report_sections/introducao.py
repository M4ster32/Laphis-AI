"""Chapter 1 — Introdução."""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...docx_utils.blocks import add_heading, add_paragraph, add_bullet


def render(doc):
    """Render chapter 1 with its four subsections."""
    add_heading(doc, "Introdução", level=1, number="1")

    add_heading(doc, "Enquadramento", level=2, number="1.1")
    add_paragraph(doc,
        "A adoção massiva de smartphones, dispositivos vestíveis (wearables) "
        "e aplicações móveis de autocuidado democratizou o acesso ao "
        "acompanhamento pessoal de indicadores de saúde. O utilizador médio "
        "regista, hoje, atividade física, ingestão calórica, padrões de sono "
        "e sinais biométricos através de uma combinação de ferramentas "
        "digitais. Contudo, o ecossistema que emergiu é profundamente "
        "fragmentado: cada aplicação domina um silo funcional (treino, "
        "nutrição, hidratação, mindfulness) e os dados raramente cruzam "
        "fronteiras, resultando em recomendações genéricas que ignoram o "
        "contexto integral do utilizador.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "Paralelamente, a evolução recente dos Modelos de Linguagem de "
        "Grande Escala (LLM) e das técnicas de Retrieval-Augmented "
        "Generation (RAG) abriu espaço para que sistemas personalizados "
        "cruzem dados biométricos, histórico individual e literatura "
        "científica, oferecendo aconselhamento adaptativo com um nível de "
        "granularidade que até recentemente não estava ao alcance de "
        "soluções de mercado [OpAI24, CiAl24].",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "É neste enquadramento que se insere o LAPHIS (Life and Physical "
        "Health Intelligent System): uma plataforma web integrada que "
        "agrega tracking multi-domínio, um motor de recomendação próprio "
        "baseado em heurísticas clínicas reconhecidas (equação de "
        "Harris-Benedict para cálculo de TDEE, fórmula BMI, distribuição "
        "de macronutrientes), e uma camada opcional de IA conversacional "
        "para aconselhamento personalizado.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Motivação e Identificação do Problema", level=2,
                number="1.2")
    add_paragraph(doc,
        "A motivação para o desenvolvimento do LAPHIS parte da observação "
        "— pessoal e documentada em estudos de utilização — de que o "
        "utilizador típico interage diariamente com entre três e cinco "
        "aplicações distintas para gerir a sua saúde e fitness. O custo "
        "cognitivo desta dispersão, associado à ausência de uma visão "
        "holística do perfil, traduz-se em abandono de hábitos e em "
        "resultados aquém do potencial.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "Três problemas concretos foram identificados como guia do "
        "trabalho:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    add_bullet(doc,
        "Fragmentação de dados — treino, nutrição e outros registos são "
        "mantidos em silos, impedindo o cruzamento necessário para "
        "recomendações integradas.")
    add_bullet(doc,
        "Recomendações estáticas e genéricas — planos tabelados que "
        "ignoram objetivos individuais, restrições alimentares, nível real "
        "de atividade e feedback contínuo.")
    add_bullet(doc,
        "Ausência de adaptação — planos que não reagem à evolução real do "
        "utilizador, à estagnação, ao progresso, ou à mudança de "
        "prioridades.")

    add_heading(doc, "Objetivos", level=2, number="1.3")
    add_paragraph(doc,
        "O objetivo geral do trabalho é a conceção, implementação e "
        "disponibilização em ambiente produtivo de uma plataforma web "
        "completa de saúde e fitness com IA, que resolva os problemas de "
        "fragmentação, personalização e adaptação identificados. Os "
        "objetivos específicos dividem-se nos seguintes pontos:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    add_bullet(doc, "Desenvolver uma aplicação web cliente-servidor com "
                    "separação clara entre SPA React e API REST FastAPI.")
    add_bullet(doc, "Implementar um motor de recomendação próprio, "
                    "suportado por heurísticas clínicas reconhecidas.")
    add_bullet(doc, "Integrar uma camada opcional de IA (OpenAI) com "
                    "pipeline RAG sobre PDFs, com fallback para modo "
                    "baseado em regras.")
    add_bullet(doc, "Garantir autenticação segura com verificação de "
                    "email, hashing bcrypt e JWT stateless.")
    add_bullet(doc, "Assegurar qualidade com suite de testes automatizados "
                    "e código livre de dependências deprecated.")
    add_bullet(doc, "Colocar o sistema em produção através de pipeline "
                    "CI/CD automatizado (Vercel + Render).")

    add_heading(doc, "Estrutura do Documento", level=2, number="1.4")
    add_paragraph(doc,
        "O presente relatório segue a estrutura recomendada para o "
        "Trabalho Final de Curso e está organizado da seguinte forma:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    add_bullet(doc, "Na Secção 2 é discutida a pertinência e viabilidade "
                    "do trabalho, incluindo análise comparativa com "
                    "soluções existentes.")
    add_bullet(doc, "Na Secção 3 apresentam-se a especificação e "
                    "modelação — requisitos funcionais e não-funcionais, "
                    "casos de uso, diagrama entidade-relação e protótipos "
                    "de interface.")
    add_bullet(doc, "Na Secção 4 descreve-se a solução desenvolvida: "
                    "arquitetura, stack tecnológico, componentes e "
                    "interfaces.")
    add_bullet(doc, "Na Secção 5 detalha-se o plano de testes e "
                    "validação.")
    add_bullet(doc, "Na Secção 6 apresenta-se o método e planeamento, com "
                    "cronograma.")
    add_bullet(doc, "Na Secção 7 documentam-se os resultados preliminares "
                    "dos testes e o grau de cumprimento dos requisitos.")
    add_bullet(doc, "Na Secção 8 conclui-se o relatório, identificando os "
                    "trabalhos futuros previstos para a entrega final.")
    add_bullet(doc, "No Anexo 1 incluem-se elementos de formatação e "
                    "referência.")

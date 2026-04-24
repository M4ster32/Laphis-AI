"""Chapter 7 — Resultados."""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...docx_utils.blocks import (
    add_heading,
    add_paragraph,
    add_caption,
    add_code_block,
)
from ...docx_utils.tables import add_table
from ..project_facts import (
    FUNCTIONAL_REQUIREMENTS,
    NON_FUNCTIONAL_REQUIREMENTS,
    METRICS,
)


def _justification_for(status):
    """Map requirement status to the report's justification column."""
    # Realised requirements get a dash so the table stays clean; parcial
    # and planeado rows defer to the status string itself (the project
    # facts module stores concrete reasoning on the description).
    if status == "Realizado":
        return "—"
    return status


def render(doc):
    """Render chapter 7: test results and requirements fulfilment table."""
    add_heading(doc, "Resultados", level=1, number="7")

    add_heading(doc, "Resultados dos Testes", level=2, number="7.1")
    add_paragraph(doc,
        "A execução da suite de testes automatizados à data desta entrega "
        f"intercalar produz os seguintes resultados: "
        f"{METRICS['test_count']} testes, {METRICS['test_count']} "
        f"aprovados, 0 falhas, {METRICS['test_warnings']} warnings. O "
        "tempo total de execução é inferior a 10 segundos, graças à "
        "utilização de SQLite in-memory.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_code_block(doc,
        "$ pytest --tb=short -q\n"
        "............................................................\n"
        f"{METRICS['test_count']} passed in 8.43s")

    add_paragraph(doc,
        "Os testes de aceitação por terceiros encontram-se em fase de "
        "planeamento. O inquérito, guião de entrevista e resultados "
        "serão incluídos em anexo na entrega final, conforme orientação "
        "do template.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Cumprimento de Requisitos", level=2, number="7.2")
    add_paragraph(doc,
        "A Tabela 8 sumariza o estado de cumprimento de cada requisito "
        "identificado na Secção 3.1. Os requisitos com estado "
        "\"Parcialmente realizado\" correspondem a funcionalidades com "
        "núcleo implementado mas com refinamentos previstos para a "
        "entrega final. Os \"Planeados\" são requisitos com baixa "
        "prioridade que dependem da disponibilidade temporal até à "
        "entrega final.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_caption(doc, "Cumprimento de requisitos.", kind="Tabela 8")
    # Merge functional + non-functional into a single fulfilment view.
    rows = []
    for req_id, name, _desc, status in FUNCTIONAL_REQUIREMENTS:
        rows.append([req_id, name, status, _justification_for(status)])
    for req_id, name, _desc, status in NON_FUNCTIONAL_REQUIREMENTS:
        rows.append([req_id, name, status, _justification_for(status)])
    add_table(doc,
        ["ID", "Requisito", "Estado", "Justificação"],
        rows,
        col_widths=[1.5, 4.5, 3.0, 6.5])

"""Chapter 8 — Conclusão."""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...docx_utils.blocks import add_heading, add_paragraph, add_bullet
from ..project_facts import METRICS


def render(doc):
    """Render chapter 8 with conclusions and future work."""
    add_heading(doc, "Conclusão", level=1, number="8")

    add_heading(doc, "Conclusão", level=2, number="8.1")
    add_paragraph(doc,
        "O presente relatório intercalar documenta o estado de "
        "desenvolvimento do LAPHIS à data de abril de 2026, cobrindo os "
        "trabalhos desde o início da UC de Projeto I em setembro de "
        "2025. O problema identificado — a fragmentação das ferramentas "
        "de tracking de saúde e a ausência de recomendações "
        "verdadeiramente personalizadas — foi abordado através da "
        "conceção e implementação de uma plataforma web integrada, que "
        "combina tracking multi-domínio, um motor de recomendação "
        "próprio baseado em heurísticas clínicas reconhecidas, e uma "
        "camada opcional de IA conversacional com pipeline RAG.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "Os principais resultados desta fase incluem: (i) 13 dos 15 "
        "requisitos funcionais encontram-se realizados, os restantes em "
        "estado parcial ou planeado; (ii) 9 dos 10 requisitos "
        f"não-funcionais cumpridos, um parcial; (iii) "
        f"{METRICS['test_count']} testes automatizados a passar com "
        f"{METRICS['test_warnings']} warnings; (iv) aplicação em fase de "
        "deploy em produção nos serviços Vercel e Render; "
        f"(v) aproximadamente {METRICS['total_lines']} linhas de código "
        f"distribuídas entre backend ({METRICS['backend_lines']} linhas) "
        f"e frontend ({METRICS['frontend_lines']} linhas); (vi) "
        "integração bem-sucedida da camada de IA com mecanismo de "
        "fallback para modo baseado em regras.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "As principais conclusões das várias fases do desenvolvimento "
        "são: a escolha do stack tecnológico (FastAPI + React + "
        "SQLAlchemy 2.0) revelou-se acertada, proporcionando alta "
        "produtividade e um ciclo de desenvolvimento rápido; a "
        "arquitetura modular facilitou significativamente a expansão "
        "incremental de funcionalidades; a separação clara entre motor "
        "de regras e camada de IA provou-se crítica para garantir o "
        "funcionamento do sistema em ambientes sem chave OpenAI; o "
        "investimento em testes automatizados desde cedo preveniu "
        "diversos bugs de regressão durante a modernização do código.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Trabalhos Futuros", level=2, number="8.2")
    add_paragraph(doc,
        "Até à entrega final de Projeto II, prevista para maio de 2026, "
        "serão realizados os seguintes trabalhos:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    for item in [
        "Conclusão do requisito RF08 — refinamento dos triggers de "
        "adaptação automática do plano diário.",
        "Conclusão do requisito RF13 — finalização da interface de "
        "upload de PDFs e consulta em linguagem natural.",
        "Implementação do requisito RF14 — notificações push via "
        "Service Worker.",
        "Aplicação do inquérito de validação a 10-15 utilizadores-teste "
        "e inclusão dos resultados em anexo.",
        "Inclusão de screenshots reais dos ecrãs e diagramas UML/ER "
        "exportados de ferramentas externas.",
        "Gravação do vídeo demonstrativo e publicação do link no "
        "repositório.",
        "Preparação de credenciais de demonstração e documentação de "
        "acesso para os avaliadores.",
    ]:
        add_bullet(doc, item)

    add_paragraph(doc,
        "Para além do TFC, identificam-se as seguintes direções de "
        "evolução futura, já fora do âmbito académico:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    for item in [
        "Integração com wearables — Apple Health, Fitbit, Google Fit via "
        "APIs oficiais (RF15).",
        "Gamificação — sistema de pontos, badges e streaks para aumentar "
        "retenção e adesão.",
        "Progressive Web App — modo offline completo com Service Worker "
        "e sincronização diferida.",
        "Machine Learning — substituição progressiva das heurísticas por "
        "modelos treinados nos dados reais dos utilizadores.",
        "Aplicação móvel nativa — port para React Native com partilha do "
        "código de lógica de negócio.",
        "Marketplace profissional — permitir a nutricionistas e personal "
        "trainers certificados publicarem planos.",
        "Scan de código de barras para dados nutricionais automáticos "
        "via API Open Food Facts.",
        "Internacionalização completa (i18n) — suporte para inglês e "
        "espanhol no mínimo.",
    ]:
        add_bullet(doc, item)

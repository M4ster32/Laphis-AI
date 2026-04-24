"""
Cover, copyright, acknowledgements, abstract (PT/EN), TOC placeholders,
figure list, table list and acronym list.

These pre-body sections share a 'no numbering' convention and are
grouped together in one module.
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...docx_utils.blocks import add_paragraph, page_break
from ...docx_utils.theme import NAVY, TEXT, GREY, set_run
from ...docx_utils.toc import add_toc_entries, add_siglas
from ..project_facts import PROJECT, SIGLAS


def _centered(doc, text, size=11, bold=False, italic=False, color=TEXT,
              name="Calibri"):
    """Utility: append a centred paragraph with custom run settings."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, name=name, size=size, bold=bold, italic=italic, color=color)
    return p


def render_cover(doc):
    """Lusofona-style cover page."""
    for _ in range(3):
        doc.add_paragraph()

    # University wordmark placeholder — real logo is inserted manually
    # in Word before final delivery.
    _centered(doc, "UNIVERSIDADE", size=16, color=NAVY,
              name="Times New Roman")
    _centered(doc, "LUSÓFONA", size=28, bold=True, color=NAVY,
              name="Times New Roman")

    for _ in range(3):
        doc.add_paragraph()

    _centered(doc, PROJECT["name"], size=36, bold=True,
              name="Times New Roman")
    _centered(doc, PROJECT["tagline"], size=16,
              name="Times New Roman")

    for _ in range(2):
        doc.add_paragraph()

    _centered(doc, "Trabalho Final de Curso", size=14, bold=True,
              name="Times New Roman")
    _centered(doc, PROJECT["degree"], size=11, bold=True)
    _centered(doc, PROJECT["course"], size=11, bold=True)
    _centered(doc, PROJECT["delivery"], size=11)
    _centered(doc, "2º Semestre", size=11)

    for _ in range(2):
        doc.add_paragraph()

    _centered(doc, PROJECT["author"], size=11, bold=True)
    doc.add_paragraph()

    # Orientador / Entidade externa — kept as placeholders.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Orientador: ")
    set_run(r, size=11, bold=True)
    r = p.add_run("[nome do orientador]")
    set_run(r, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Entidade Externa: ")
    set_run(r, size=11, bold=True)
    r = p.add_run("<se aplicável>")
    set_run(r, size=11)

    for _ in range(2):
        doc.add_paragraph()

    _centered(doc, PROJECT["department"], size=10.5)
    _centered(doc, f"{PROJECT['university']}, {PROJECT['faculty']}", size=10.5)
    _centered(doc, PROJECT["date"], size=10.5)
    doc.add_paragraph()
    _centered(doc, "w w w . u l u s o f o n a . p t", size=9, color=NAVY)

    page_break(doc)


def render_copyright(doc):
    """Direitos de cópia — official Lusofona wording."""
    for _ in range(10):
        doc.add_paragraph()

    add_paragraph(doc, "Direitos de cópia", size=14, bold=True,
                  color=NAVY, space_after=12)
    add_paragraph(doc,
        f"{PROJECT['name']} — {PROJECT['tagline']}, Copyright de "
        f"{PROJECT['author']}, {PROJECT['university']}.",
        size=11, space_after=10)
    add_paragraph(doc,
        "A Faculdade de Ciências Naturais, Engenharia e Tecnologias (FCNET) "
        "e a Universidade Lusófona têm o direito, perpétuo e sem limites "
        "geográficos, de arquivar e publicar esta dissertação/monografia "
        "através de exemplares impressos reproduzidos em papel ou de forma "
        "digital, ou por qualquer outro meio conhecido ou que venha a ser "
        "inventado, e de a divulgar através de repositórios científicos e "
        "de admitir a sua cópia e distribuição com objetivos educacionais "
        "ou de investigação, não comerciais, desde que seja dado crédito "
        "ao autor e editor.",
        size=11)

    page_break(doc)


def render_acknowledgements(doc):
    """Agradecimentos."""
    add_paragraph(doc, "Agradecimentos", size=16, bold=True, color=NAVY,
                  space_after=18)
    add_paragraph(doc,
        "Aos professores do Departamento de Engenharia Informática e "
        "Sistemas de Informação da Universidade Lusófona, pelo "
        "acompanhamento ao longo das unidades curriculares de Projeto I "
        "e Projeto II.", size=11, space_after=8)
    add_paragraph(doc,
        "Ao orientador deste trabalho, pela disponibilidade, sugestões e "
        "apreciação crítica nas várias fases do desenvolvimento do projeto.",
        size=11, space_after=8)
    add_paragraph(doc,
        "Aos colegas que testaram versões preliminares da aplicação e "
        "cujo feedback contribuiu para decisões de desenho e "
        "funcionalidade.", size=11, space_after=8)
    add_paragraph(doc,
        "À família e amigos, pelo apoio durante o percurso académico que "
        "culmina neste relatório.", size=11)

    page_break(doc)


def render_resumo(doc):
    """Resumo PT — one single paragraph, as per template."""
    add_paragraph(doc, "Resumo", size=16, bold=True, color=NAVY, space_after=12)
    add_paragraph(doc,
        "O LAPHIS (Life and Physical Health Intelligent System) é uma "
        "plataforma web de gestão integrada de saúde e fitness, "
        "desenvolvida no âmbito do Trabalho Final de Curso da Licenciatura "
        "em Engenharia Informática. O problema em estudo é a fragmentação "
        "das ferramentas atuais de tracking pessoal — treino, nutrição, "
        "hidratação, peso e bem-estar mental encontram-se dispersos por "
        "aplicações distintas, sem cruzamento de dados nem recomendações "
        "adaptadas ao perfil individual. Para resolver este problema, foi "
        "desenvolvida uma aplicação cliente-servidor com frontend em "
        "React 19 (Vite) e backend em FastAPI (Python 3.11), assente em "
        "base de dados relacional gerida por SQLAlchemy 2.0. A "
        "inteligência do sistema é suportada por um motor de recomendação "
        "próprio de 857 linhas, baseado em heurísticas nutricionais e de "
        "exercício (BMI, TDEE, distribuição de macronutrientes, plano "
        "diário adaptativo), complementado por uma camada opcional de "
        "Inteligência Artificial (OpenAI GPT) com pipeline de "
        "Retrieval-Augmented Generation sobre documentos PDF. O sistema "
        "foi validado por uma suite automatizada de 62 testes (pytest + "
        "httpx) e encontra-se atualmente em fase de deploy em produção "
        "(Vercel para o frontend, Render para o backend e base de dados "
        "PostgreSQL). O relatório documenta a arquitetura, tecnologias, "
        "especificação de requisitos, decisões de desenho, progresso até "
        "à entrega intercalar e trabalhos futuros previstos para a entrega "
        "final.",
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=14)

    p = doc.add_paragraph()
    r = p.add_run("Palavras-chave: ")
    set_run(r, size=11, bold=True)
    r = p.add_run("saúde digital; inteligência artificial; recomendação "
                  "personalizada; aplicação web; FastAPI")
    set_run(r, size=11)

    page_break(doc)


def render_abstract(doc):
    """English abstract, mirrors the PT resumo."""
    add_paragraph(doc, "Abstract", size=16, bold=True, color=NAVY,
                  space_after=12)
    add_paragraph(doc,
        "LAPHIS (Life and Physical Health Intelligent System) is an "
        "integrated web platform for health and fitness management, "
        "developed within the scope of the Final Course Project for the "
        "Bachelor's degree in Computer Engineering. The problem addressed "
        "is the fragmentation of current personal tracking tools — "
        "workouts, nutrition, hydration, weight and mental well-being "
        "are scattered across separate applications, without cross-domain "
        "data correlation or truly personalized recommendations. To tackle "
        "this, a client-server application was developed with a React 19 "
        "(Vite) frontend and a FastAPI (Python 3.11) backend, built on a "
        "relational database managed through SQLAlchemy 2.0. Intelligence "
        "is provided by a custom 857-line recommendation engine based on "
        "nutritional and exercise heuristics (BMI, TDEE, macronutrient "
        "distribution, adaptive daily plan), complemented by an optional "
        "Artificial Intelligence layer (OpenAI GPT) with a "
        "Retrieval-Augmented Generation pipeline over PDF documents. The "
        "system is validated by an automated suite of 62 tests (pytest + "
        "httpx) and is currently being deployed to production (Vercel for "
        "the frontend, Render for the backend and PostgreSQL database). "
        "This report documents the architecture, technology stack, "
        "requirements specification, design decisions, progress up to the "
        "intercalar delivery and planned future work for the final "
        "delivery.",
        size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=14)

    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    set_run(r, size=11, bold=True)
    r = p.add_run("digital health; artificial intelligence; personalized "
                  "recommendation; web application; FastAPI")
    set_run(r, size=11)

    page_break(doc)


def render_toc(doc):
    """Manual TOC placeholder — user should replace with auto-TOC in Word."""
    add_paragraph(doc, "Índice", size=16, bold=True, color=NAVY, space_after=12)
    add_paragraph(doc,
        "[Este índice deve ser atualizado automaticamente no Word usando "
        "Referências → Índice. Clicar com o botão direito → Atualizar "
        "campos.]", size=10, italic=True, color=GREY, space_after=12)

    toc_entries = [
        ("Agradecimentos", "iii"),
        ("Resumo", "iv"),
        ("Abstract", "v"),
        ("Índice", "vi"),
        ("Lista de Figuras", "viii"),
        ("Lista de Tabelas", "ix"),
        ("Lista de Siglas", "x"),
        ("1  Introdução", "1"),
        ("    1.1  Enquadramento", "1"),
        ("    1.2  Motivação e Identificação do Problema", "2"),
        ("    1.3  Objetivos", "3"),
        ("    1.4  Estrutura do Documento", "3"),
        ("2  Pertinência e Viabilidade", "4"),
        ("    2.1  Pertinência", "4"),
        ("    2.2  Viabilidade", "5"),
        ("    2.3  Análise Comparativa com Soluções Existentes", "5"),
        ("    2.4  Proposta de Inovação e Mais-Valias", "7"),
        ("    2.5  Identificação de Oportunidade de Negócio", "7"),
        ("3  Especificação e Modelação", "8"),
        ("    3.1  Análise de Requisitos", "8"),
        ("    3.2  Modelação", "11"),
        ("    3.3  Protótipos de Interface", "13"),
        ("4  Solução Desenvolvida", "14"),
        ("    4.1  Introdução", "14"),
        ("    4.2  Arquitetura", "14"),
        ("    4.3  Tecnologias e Ferramentas Utilizadas", "16"),
        ("    4.4  Ambientes de Teste e de Produção", "17"),
        ("    4.5  Abrangência", "18"),
        ("    4.6  Componentes", "18"),
        ("    4.7  Interfaces", "20"),
        ("5  Testes e Validação", "22"),
        ("6  Método e Planeamento", "24"),
        ("7  Resultados", "26"),
        ("    7.1  Resultados dos Testes", "26"),
        ("    7.2  Cumprimento de Requisitos", "27"),
        ("8  Conclusão", "28"),
        ("    8.1  Conclusão", "28"),
        ("    8.2  Trabalhos Futuros", "29"),
        ("Bibliografia", "30"),
        ("Anexo 1 — Recomendações para formatação de um relatório", "31"),
    ]
    add_toc_entries(doc, toc_entries)
    page_break(doc)


def render_figure_list(doc):
    """Lista de Figuras."""
    add_paragraph(doc, "Lista de Figuras", size=16, bold=True, color=NAVY,
                  space_after=12)
    figures = [
        ("Figura 1 – Arquitetura geral do sistema LAPHIS.", "15"),
        ("Figura 2 – Diagrama Entidade-Relação da base de dados.", "11"),
        ("Figura 3 – Fluxo de autenticação e verificação de email.", "19"),
        ("Figura 4 – Pipeline de RAG sobre documentos PDF.", "20"),
        ("Figura 5 – Mapa aplicacional e navegação entre ecrãs.", "13"),
        ("Figura 6 – Diagrama de casos de uso (ator utilizador).", "10"),
        ("Figura 7 – Cronograma Gantt do projeto.", "25"),
    ]
    add_toc_entries(doc, figures)
    page_break(doc)


def render_table_list(doc):
    """Lista de Tabelas."""
    add_paragraph(doc, "Lista de Tabelas", size=16, bold=True, color=NAVY,
                  space_after=12)
    tables_list = [
        ("Tabela 1 – Análise de benchmarking comparativo com soluções "
         "existentes.", "6"),
        ("Tabela 2 – Requisitos funcionais identificados.", "8"),
        ("Tabela 3 – Requisitos não-funcionais identificados.", "10"),
        ("Tabela 4 – Entidades do modelo de dados.", "12"),
        ("Tabela 5 – Tecnologias utilizadas no backend.", "16"),
        ("Tabela 6 – Tecnologias utilizadas no frontend.", "16"),
        ("Tabela 7 – Plano de testes por módulo.", "22"),
        ("Tabela 8 – Cumprimento de requisitos.", "27"),
    ]
    add_toc_entries(doc, tables_list)
    page_break(doc)


def render_siglas(doc):
    """Lista de Siglas — pulls from project_facts.SIGLAS."""
    add_paragraph(doc, "Lista de Siglas", size=16, bold=True, color=NAVY,
                  space_after=12)
    add_siglas(doc, SIGLAS)


def render(doc):
    """
    Render all front-matter sections in canonical template order.

    :param doc: python-docx Document.
    """
    render_cover(doc)
    render_copyright(doc)
    render_acknowledgements(doc)
    render_resumo(doc)
    render_abstract(doc)
    render_toc(doc)
    render_figure_list(doc)
    render_table_list(doc)
    render_siglas(doc)

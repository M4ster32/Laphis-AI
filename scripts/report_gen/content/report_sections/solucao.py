"""Chapter 4 — Solução Desenvolvida."""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...docx_utils.blocks import (
    add_heading,
    add_paragraph,
    add_bullet,
    add_caption,
    add_code_block,
)
from ...docx_utils.tables import add_table
from ...docx_utils.theme import GREY
from ..project_facts import (
    PROJECT,
    BACKEND_STACK,
    FRONTEND_STACK,
    BACKEND_ROUTERS,
    FRONTEND_COMPONENTS,
    METRICS,
)


def render(doc):
    """Render chapter 4 with architecture diagram and stack tables."""
    add_heading(doc, "Solução Desenvolvida", level=1, number="4")

    add_heading(doc, "Introdução", level=2, number="4.1")
    add_paragraph(doc,
        "A solução desenvolvida, LAPHIS, é uma aplicação web de "
        "arquitetura cliente-servidor, com frontend Single-Page "
        "Application (SPA) em React 19 e backend REST em FastAPI. A "
        "comunicação é feita via HTTP/JSON com autenticação JWT "
        "stateless; o backend comunica com uma base de dados relacional "
        "(SQLite em desenvolvimento, PostgreSQL em produção) através do "
        "ORM SQLAlchemy 2.0.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "Face às soluções apresentadas no benchmarking (Secção 2.3), o "
        "LAPHIS distingue-se pela integração de múltiplos domínios num "
        "único produto, pela existência de um motor de recomendação "
        "baseado em heurísticas clínicas transparentes, e pela camada de "
        "IA opcional com pipeline RAG.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

    add_paragraph(doc, "Artefactos da entrega final (a completar):",
                  bold=True, space_after=4)
    add_bullet(doc, "Link para vídeo demonstrativo: [a adicionar antes da "
                    "entrega final].")
    add_bullet(doc, f"Repositório Git público: {PROJECT['repo']}")
    add_bullet(doc,
        f"Solução funcional em produção: {PROJECT['frontend_url']} "
        "(credenciais de demonstração a fornecer antes da entrega final).")

    add_paragraph(doc,
        "As próximas subsecções detalham a arquitetura, o stack "
        "tecnológico, os ambientes de execução, a abrangência curricular "
        "do trabalho, os componentes implementados e os principais ecrãs "
        "da interface.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Arquitetura", level=2, number="4.2")
    add_paragraph(doc,
        "A arquitetura do LAPHIS divide-se em quatro camadas principais, "
        "representadas na Figura 1. A separação entre camadas foi "
        "orientada por princípios de responsabilidade única, stateless "
        "por omissão (exceto na persistência) e independência tecnológica "
        "— o frontend e o backend comunicam exclusivamente por API REST, "
        "permitindo substituir qualquer das camadas sem afetar as "
        "restantes.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_code_block(doc,
        "┌──────────────┐   HTTPS   ┌──────────────┐   TCP    ┌──────────────┐\n"
        "│   Browser    │ ────────▶ │    Vercel    │ ───────▶ │    Render    │\n"
        "│  React 19    │  JSON/JWT │  Edge CDN    │          │ FastAPI app  │\n"
        "└──────────────┘           └──────────────┘          └──────┬───────┘\n"
        "                                                            │\n"
        "                                                     ┌──────▼───────┐\n"
        "                                                     │  PostgreSQL  │\n"
        "                                                     │  (Render)    │\n"
        "                                                     └──────────────┘")

    add_caption(doc, "Arquitetura geral do sistema LAPHIS.",
                kind="Figura 1")

    add_paragraph(doc,
        "As opções arquiteturais foram fundamentadas nos seguintes "
        "pontos:", space_after=4)
    add_bullet(doc, "SPA em vez de server-rendered pages: permite "
                    "interações ricas e fluídas, reduz roundtrips e "
                    "simplifica o deploy separado (frontend estático em "
                    "CDN).")
    add_bullet(doc, "REST em vez de GraphQL: a complexidade do domínio "
                    "é moderada e REST tem menor overhead de ferramentas "
                    "para um projeto académico.")
    add_bullet(doc, "JWT stateless em vez de sessões em cookies "
                    "server-side: facilita o escalamento horizontal e "
                    "remove a necessidade de store de sessões.")
    add_bullet(doc, "SQLAlchemy com SQLite em dev e PostgreSQL em prod: "
                    "o ORM abstrai a diferença, permitindo ciclo rápido "
                    "de desenvolvimento.")

    add_heading(doc, "Tecnologias e Ferramentas Utilizadas",
                level=2, number="4.3")
    add_paragraph(doc,
        "A Tabela 5 e a Tabela 6 listam as tecnologias principais de "
        "backend e frontend, respetivamente, com justificação do uso.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_caption(doc, "Tecnologias utilizadas no backend.", kind="Tabela 5")
    add_table(doc,
        ["Tecnologia", "Versão", "Justificação"],
        [list(row) for row in BACKEND_STACK],
        col_widths=[3.5, 2.0, 10.5])

    add_caption(doc, "Tecnologias utilizadas no frontend.", kind="Tabela 6")
    add_table(doc,
        ["Tecnologia", "Versão", "Justificação"],
        [list(row) for row in FRONTEND_STACK],
        col_widths=[3.5, 2.0, 10.5])

    add_heading(doc, "Ambientes de Teste e de Produção", level=2,
                number="4.4")
    add_paragraph(doc,
        "O projeto utiliza três ambientes distintos, com configuração "
        "controlada por variáveis de ambiente:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    add_bullet(doc, "Desenvolvimento local — Python virtualenv, base de "
                    "dados SQLite em ficheiro local, frontend servido "
                    "pelo Vite dev server em http://localhost:5173.")
    add_bullet(doc, "Testes automatizados — base de dados SQLite "
                    "in-memory, cliente HTTP httpx assíncrono, fixtures "
                    "pytest com isolamento total por teste.")
    add_bullet(doc, "Produção — frontend deployado na Vercel com CDN "
                    "global, backend na Render com 1 vCPU / 512 MB RAM, "
                    "base de dados PostgreSQL gerida (tier gratuito "
                    "Render, 1 GB de armazenamento).")
    add_paragraph(doc,
        "Os recursos atuais são suficientes para a fase de demonstração "
        "e validação por terceiros. Num cenário de exploração comercial, "
        "recomenda-se escalar o backend para tier pago (2 vCPU, 2 GB "
        "RAM) e migrar a base de dados para plano Standard (10 GB, "
        "backups diários).",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Abrangência", level=2, number="4.5")
    add_paragraph(doc,
        "O desenvolvimento do LAPHIS aplica e integra conteúdos das "
        "seguintes unidades curriculares da Licenciatura em Engenharia "
        "Informática:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    for uc in [
        "Programação Web — React, CSS, responsividade, consumo de APIs.",
        "Bases de Dados — modelação relacional, normalização, SQL.",
        "Engenharia de Software — requisitos, UML, testes, ciclo de vida.",
        "Análise de Sistemas — levantamento de requisitos, casos de uso, "
        "benchmarking.",
        "Sistemas Distribuídos — arquitetura cliente-servidor, REST, "
        "autenticação stateless.",
        "Segurança Informática — hashing bcrypt, JWT, HTTPS, CORS.",
        "Inteligência Artificial — integração de LLM, pipeline RAG, "
        "heurísticas de recomendação.",
    ]:
        add_bullet(doc, uc)

    add_heading(doc, "Componentes", level=2, number="4.6")

    add_heading(doc, "Backend — API REST com FastAPI", level=3,
                number="4.6.1")
    add_paragraph(doc,
        f"O backend organiza-se em {METRICS['router_count']} routers "
        "FastAPI, cada um responsável por um domínio funcional. O ponto "
        "de entrada é o ficheiro main.py, que regista todos os routers, "
        "configura CORS, e utiliza um lifespan handler para inicialização "
        "da base de dados e migrações leves.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Compose the routers summary line from project_facts, so a change in
    # the actual codebase only requires an update in project_facts.
    router_summary = ", ".join(
        f"{name} ({count} endpoints)" for name, count, _ in BACKEND_ROUTERS
    )
    add_paragraph(doc,
        f"Os routers implementados são: {router_summary}, totalizando "
        f"{METRICS['endpoint_total']} endpoints.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "O motor de recomendação (core/recommender.py, "
        f"{METRICS['recommender_lines']} linhas) implementa as "
        "heurísticas clínicas reconhecidas: BMI (peso/altura²), TDEE "
        "(equação de Harris-Benedict revista multiplicada pelo fator de "
        "atividade), distribuição de macronutrientes por objetivo, "
        "hidratação recomendada (30-35 ml/kg) e planeamento diário.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "A camada de IA (core/ai_engine.py e core/rag.py) integra a API "
        "OpenAI. A deteção da variável de ambiente OPENAI_API_KEY é "
        "feita no arranque: se presente, o modo IA é ativado; caso "
        "contrário, o sistema opera em modo baseado em regras sem "
        "qualquer falha. O pipeline RAG extrai texto de PDFs via "
        "PyMuPDF, segmenta em chunks e injeta o contexto relevante no "
        "prompt enviado ao LLM.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Frontend — SPA React", level=3, number="4.6.2")
    add_paragraph(doc,
        f"O frontend organiza-se em {METRICS['pages_count']} páginas e "
        f"{METRICS['component_count']} componentes reutilizáveis. O "
        "routing, gerido pelo React Router 7, distingue rotas públicas "
        "(landing, autenticação) de rotas protegidas (toda a área /app), "
        "estas últimas exigindo token JWT válido no contexto.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "A gestão de estado é feita via React Context API em dois "
        "contextos globais: AppContext (token, perfil do utilizador, "
        "flags de loading) e ThemeContext (modo dark/light, persistência "
        "em localStorage). A comunicação HTTP é centralizada em "
        "services/api.js, um wrapper que injeta automaticamente o header "
        "Authorization: Bearer <token> em todos os pedidos autenticados.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    components_line = ", ".join(FRONTEND_COMPONENTS)
    add_paragraph(doc,
        f"Os componentes reutilizáveis incluem: {components_line}.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Interfaces", level=2, number="4.7")
    add_paragraph(doc,
        "Os ecrãs mais representativos da aplicação — a incluir aqui "
        "como screenshots na entrega final — são:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    for desc in [
        "Dashboard — visão geral do dia com KPIs de calorias, macros, "
        "água e peso; plano do dia abaixo.",
        "Logs — registo unificado de treinos e refeições, com filtros "
        "por data e categoria.",
        "Plans / PlanDetail — lista de planos ativos e detalhe editável "
        "com geração por IA ou regras.",
        "Chat — interface conversacional com bubbles, histórico por "
        "sessão e input com suporte a multi-linha.",
        "Reports — gráficos Recharts (linhas, barras, radial) e botão "
        "de exportação PDF.",
        "Zen — seleção de tipo de meditação, timer circular e registo "
        "de humor antes/depois.",
    ]:
        add_bullet(doc, desc)

    add_paragraph(doc,
        "[Espaço reservado para screenshots dos ecrãs principais. A "
        "incluir na entrega final com capturas em modo escuro em "
        "resolução 1440×900.]",
        italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

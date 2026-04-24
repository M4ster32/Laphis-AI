"""Bibliography and Anexo 1 — trailing sections of the report."""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from ...docx_utils.blocks import add_heading, add_paragraph, add_bullet
from ...docx_utils.theme import set_run
from ..project_facts import REFERENCES


def render_bibliography(doc):
    """Bibliography block with [AaAa20]-style keys."""
    # Level-1 heading with no number, as per template.
    add_heading(doc, "Bibliografia", level=1, number=None)

    for key, text in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(2.0)
        p.paragraph_format.first_line_indent = Cm(-2.0)
        run = p.add_run(key)
        set_run(run, size=10.5, bold=True)
        run = p.add_run(f"\t{text}")
        set_run(run, size=10.5)


def render_anexo_1(doc):
    """Anexo 1 — formatting recommendations."""
    add_heading(doc,
        "Anexo 1 — Recomendações para formatação do relatório",
        level=1, number=None)

    add_paragraph(doc,
        "Este anexo segue a orientação do template original da "
        "Universidade Lusófona, documentando convenções de formatação "
        "aplicadas neste relatório e servindo de referência para "
        "revisões futuras.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc, "Formatação geral", bold=True, space_after=4)
    for item in [
        "Fonte do corpo do texto: Calibri 11 pt, espaçamento entre "
        "linhas 1.15.",
        "Margens: 2.5 cm (topo, base, direita); 3.0 cm (esquerda) para "
        "encadernação.",
        "Cabeçalho: título do trabalho em itálico, alinhado à direita.",
        "Rodapé: número de página alinhado à direita.",
        "Títulos: Calibri bold, cor azul corporativo (#1F3A68), com "
        "numeração hierárquica (1, 1.1, 1.1.1).",
    ]:
        add_bullet(doc, item)

    add_paragraph(doc, "Figuras e tabelas", bold=True, space_after=4)
    for item in [
        "Numeradas sequencialmente (Figura 1, Figura 2, ...; Tabela 1, "
        "Tabela 2, ...).",
        "Legendas em Calibri 9.5 pt itálico, centradas, imediatamente "
        "abaixo da figura ou acima da tabela.",
        "Devem ser sempre referenciadas no texto antes de aparecerem "
        "(\"ver Figura 1\", \"conforme apresentado na Tabela 2\").",
    ]:
        add_bullet(doc, item)

    add_paragraph(doc, "Referências bibliográficas", bold=True,
                  space_after=4)
    for item in [
        "Formato: chave com 4 caracteres iniciais dos nomes dos autores "
        "+ 2 dígitos do ano — ex.: [TaWe20], [OpAI24].",
        "Ordenação alfabética pela chave.",
        "Citações no texto entre parêntesis retos — ex.: \"segundo "
        "[TaWe20], ...\".",
    ]:
        add_bullet(doc, item)

    add_paragraph(doc,
        "Recomenda-se a utilização das funcionalidades \"Referências → "
        "Índice\" e \"Referências → Inserir Legenda\" do MS Word para "
        "manter numeração e paginação automaticamente atualizadas. Após "
        "edições substanciais, selecionar todo o documento e premir F9 "
        "para regenerar índices.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def render(doc):
    """Render bibliography followed by Anexo 1."""
    render_bibliography(doc)
    render_anexo_1(doc)

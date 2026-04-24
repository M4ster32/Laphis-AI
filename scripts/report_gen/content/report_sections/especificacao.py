"""Chapter 3 — Especificação e Modelação."""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...docx_utils.blocks import (
    add_heading,
    add_paragraph,
    add_bullet,
    add_caption,
)
from ...docx_utils.tables import add_table
from ...docx_utils.theme import GREY
from ..project_facts import (
    FUNCTIONAL_REQUIREMENTS,
    NON_FUNCTIONAL_REQUIREMENTS,
    USE_CASES,
    ENTITIES,
)


def render(doc):
    """Render chapter 3 pulling requirements data from project_facts."""
    add_heading(doc, "Especificação e Modelação", level=1, number="3")

    add_heading(doc, "Análise de Requisitos", level=2, number="3.1")
    add_paragraph(doc,
        "A análise de requisitos resulta do cruzamento da investigação "
        "sobre soluções existentes (Secção 2.3), das especificações da UC "
        "de Análise de Sistemas e de conversas iterativas com potenciais "
        "utilizadores. A enumeração inicial, estabelecida em Projeto I, "
        "foi revista nesta entrega intercalar — os requisitos modificados "
        "são indicados explicitamente.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "Requisitos Funcionais", level=3, number="3.1.1")
    add_paragraph(doc,
        "A Tabela 2 apresenta a lista de requisitos funcionais "
        "identificados, com indicação do estado atual de implementação. "
        "O estado \"Realizado\" indica implementação completa e testada; "
        "\"Parcial\" indica funcionalidade básica existente mas com "
        "refinamentos pendentes para a entrega final.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_caption(doc, "Requisitos funcionais identificados.", kind="Tabela 2")
    add_table(doc,
        ["ID", "Requisito", "Descrição", "Estado"],
        [[r[0], r[1], r[2], r[3]] for r in FUNCTIONAL_REQUIREMENTS],
        col_widths=[1.2, 3.5, 7.5, 2.0])

    add_heading(doc, "Requisitos Não-Funcionais", level=3, number="3.1.2")
    add_paragraph(doc,
        "A Tabela 3 apresenta os requisitos não-funcionais que governam "
        "a qualidade do sistema em aspetos transversais — desempenho, "
        "segurança, usabilidade e manutenibilidade.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_caption(doc, "Requisitos não-funcionais identificados.",
                kind="Tabela 3")
    add_table(doc,
        ["ID", "Categoria", "Descrição", "Estado"],
        [[r[0], r[1], r[2], r[3]] for r in NON_FUNCTIONAL_REQUIREMENTS],
        col_widths=[1.2, 2.5, 8.5, 2.0])

    add_heading(doc, "Casos de Uso", level=3, number="3.1.3")
    add_paragraph(doc,
        "Os casos de uso foram modelados tendo como ator primário o "
        "\"Utilizador registado\" e, como atores secundários, o "
        "\"Visitante (não registado)\" e o \"Assistente IA\". A Figura 6 "
        "esquematiza os casos de uso principais.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc, "Os principais casos de uso identificados são:",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    for uc_id, uc_name in USE_CASES:
        add_bullet(doc, f"{uc_id} — {uc_name}")

    add_paragraph(doc,
        "[Espaço reservado para Figura 6 — Diagrama de casos de uso. A "
        "incluir diagrama UML elaborado em ferramenta externa "
        "(PlantUML/Draw.io).]",
        italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "Modelação", level=2, number="3.2")
    add_paragraph(doc,
        "O modelo de dados segue o paradigma relacional, com o utilizador "
        "(User) como entidade raiz. Todas as demais entidades mantêm "
        "chave estrangeira para User, com cascade configurado para "
        "permitir eliminação da conta com limpeza automática dos registos "
        "associados. Campos de data/hora são timezone-aware; estruturas "
        "semi-estruturadas (metrics, plan_data) usam campos JSON para "
        "manter flexibilidade.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "A Figura 2 apresenta o diagrama entidade-relação completo. A "
        "Tabela 4 detalha cada entidade, campos relevantes e propósito.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "[Espaço reservado para Figura 2 — Diagrama Entidade-Relação. A "
        "substituir por diagrama ER exportado do dbdiagram.io ou "
        "DrawSQL.]",
        italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_caption(doc, "Entidades do modelo de dados.", kind="Tabela 4")
    add_table(doc,
        ["Entidade", "Campos principais", "Propósito"],
        [list(row) for row in ENTITIES],
        col_widths=[3.0, 7.0, 5.0])

    add_heading(doc, "Protótipos de Interface", level=2, number="3.3")
    add_paragraph(doc,
        "O mapa aplicacional compreende 17 ecrãs organizados em duas "
        "grandes áreas: área pública (landing page, fluxos de "
        "autenticação) e área privada (sob autenticação JWT). A Figura 5 "
        "representa a navegação entre ecrãs.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
        "Os protótipos seguem um sistema de design denominado \"Premium "
        "Monochromatic\", inspirado em ferramentas como Linear, Notion e "
        "Stripe, com suporte completo a dark/light mode implementado via "
        "CSS Custom Properties. A tipografia escala-se em seis níveis "
        "(Display 36pt → Micro 12pt) e o espaçamento obedece a um grid "
        "de 4 pontos.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc,
        "[Espaço reservado para Figura 5 — Mapa aplicacional com "
        "screenshots reduzidos dos ecrãs principais. A incluir protótipos "
        "Figma exportados.]",
        italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

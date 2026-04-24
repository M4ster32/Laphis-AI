"""
Document-level setup: header, footer, page numbering.

Lives alongside theme.py because it's document chrome, not block
primitives.
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .theme import set_run, GREY


def configure_header(doc, title):
    """
    Add a right-aligned italic running title to the document header.

    Matches the Lusofona template convention (title of the work in
    italic, top-right of each page).

    :param doc: python-docx Document.
    :param title: Header text to display.
    """
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(title)
    set_run(run, size=10, italic=True, color=GREY)


def configure_footer(doc):
    """
    Add an auto-updating page number to the footer (right-aligned).

    python-docx has no high-level API for PAGE fields, so we build the
    w:fldSimple element directly.

    :param doc: python-docx Document.
    """
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Field instruction "PAGE" is the OOXML equivalent of Word's
    # auto-numbered page field.
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    size_el = OxmlElement("w:sz")
    # 20 half-points = 10pt font in OOXML units.
    size_el.set(qn("w:val"), "20")
    run_props.append(size_el)
    run.append(run_props)

    text_el = OxmlElement("w:t")
    # Placeholder "1" — Word replaces it with the actual page number
    # when the field is updated.
    text_el.text = "1"
    run.append(text_el)

    fld.append(run)
    fp._p.append(fld)

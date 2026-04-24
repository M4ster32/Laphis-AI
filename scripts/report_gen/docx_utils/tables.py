"""
Table builder with Lusofona navy header shading.

Extracted from blocks.py because table rendering has its own complexity
(header cell shading, column widths) and benefits from being isolated.
"""

from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .theme import set_run


def add_table(doc, headers, rows, col_widths=None):
    """
    Insert a Grid-styled table with a navy header row.

    :param doc: python-docx Document.
    :param headers: List of header cell strings.
    :param rows: List of row tuples/lists (strings).
    :param col_widths: Optional list of column widths in centimetres.
    :returns: The created table.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row — bold white text on navy background.
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # Cell shading via direct OOXML (python-docx has no helper).
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3A68")
        tcPr.append(shd)

    # Body rows — plain 9.5pt Calibri.
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            set_run(run, size=9.5)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    # Trailing blank paragraph for breathing room after the table.
    doc.add_paragraph()
    return table

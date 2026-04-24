"""
High-level paragraph/heading/list/code block builders.

Each function hides the repetitive python-docx plumbing (alignment,
spacing, run creation) so the content module can stay declarative.
"""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .theme import NAVY, GREY, TEXT, set_run


def add_paragraph(doc, text="", size=11, bold=False, italic=False, color=None,
                  align=None, space_after=6, space_before=0, first_line=None):
    """
    Append a styled paragraph to the document.

    :param doc: python-docx Document instance.
    :param text: Paragraph body text.
    :param size: Font size in points.
    :param bold: Bold flag.
    :param italic: Italic flag.
    :param color: Optional RGBColor; defaults to style colour.
    :param align: Optional WD_ALIGN_PARAGRAPH enum value.
    :param space_after: Trailing spacing in points.
    :param space_before: Leading spacing in points.
    :param first_line: Optional first-line indent in centimetres.
    :returns: The created paragraph.
    """
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(first_line)
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1, number=None):
    """
    Add a numbered heading with Lusofona-style formatting.

    Heading level 1 forces a page break before, matching the template's
    chapter-starts-on-new-page convention.

    :param doc: python-docx Document instance.
    :param text: Heading text.
    :param level: 1, 2 or 3.
    :param number: Optional manual number string (e.g. "1.2.3"); when
        provided, it is prepended followed by a tab.
    :returns: The created paragraph.
    """
    # Heading point sizes calibrated against the reference PDF.
    sizes = {1: 18, 2: 14, 3: 12}

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    # Prevents the heading from dangling at the bottom of a page without
    # any body content beneath it.
    p.paragraph_format.keep_with_next = True

    if number:
        run_num = p.add_run(f"{number}\t")
        set_run(run_num, size=sizes.get(level, 12), bold=True, color=NAVY)

    run = p.add_run(text)
    set_run(run, size=sizes.get(level, 12), bold=True, color=NAVY)

    if level == 1:
        p.paragraph_format.page_break_before = True
    return p


def add_bullet(doc, text, level=0, size=11):
    """
    Append a bullet-list item.

    :param doc: python-docx Document.
    :param text: Bullet text.
    :param level: Indentation level (0 = top level).
    :param size: Font size in points.
    :returns: The created paragraph.
    """
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.75)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, size=size)
    return p


def add_numbered(doc, text, size=11):
    """
    Append an auto-numbered list item.

    :param doc: python-docx Document.
    :param text: List item text.
    :param size: Font size in points.
    :returns: The created paragraph.
    """
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, size=size)
    return p


def add_caption(doc, text, kind="Figura"):
    """
    Centred italic caption, used beneath figures and above tables.

    Kind is typically 'Figura' or 'Tabela' and appears before the
    dash-separated caption text, matching the template convention.

    :param doc: python-docx Document.
    :param text: Caption body (without the 'Figura X – ' prefix).
    :param kind: Caption prefix — defaults to 'Figura'.
    :returns: The created paragraph.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"{kind} – {text}")
    set_run(run, size=9.5, italic=True, color=GREY)
    return p


def add_code_block(doc, text, size=9):
    """
    Render a monospace code block with a light grey background.

    Uses raw OOXML for cell shading because python-docx has no
    high-level API for paragraph background colour.

    :param doc: python-docx Document.
    :param text: Multi-line code text (newlines preserved).
    :param size: Monospace font size in points.
    """
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # Paragraph shading applied via w:shd element inside w:pPr —
        # python-docx exposes no helper for this.
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F3F5")
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        set_run(run, name="Consolas", size=size, color=TEXT)
    doc.add_paragraph()


def page_break(doc):
    """
    Insert a hard page break.

    :param doc: python-docx Document.
    """
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

"""
Table-of-contents-style entries and acronym list helpers.

These share a dot-leader tab stop layout so they live together.
"""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .theme import set_run


def add_toc_entries(doc, entries, right_tab_cm=15.5):
    """
    Append a list of 'label .... page' lines with a right-aligned tab.

    Used for manual TOC placeholders, figure lists and table lists.

    :param doc: python-docx Document.
    :param entries: Iterable of (label, page) tuples.
    :param right_tab_cm: Right tab stop position in centimetres.
    """
    for entry, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        tabs = p.paragraph_format.tab_stops
        # Leader 1 is the dotted leader in the OOXML tab spec.
        tabs.add_tab_stop(Cm(right_tab_cm), WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        run = p.add_run(f"{entry}\t{page}")
        set_run(run, size=10.5)


def add_siglas(doc, siglas, tab_cm=3.5):
    """
    Render the 'Lista de Siglas' two-column layout: bold acronym +
    tab-separated definition.

    :param doc: python-docx Document.
    :param siglas: Iterable of (acronym, definition) tuples.
    :param tab_cm: Tab stop for the definition column in centimetres.
    """
    for sigla, desc in siglas:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(tab_cm))
        run = p.add_run(sigla)
        set_run(run, size=10.5, bold=True)
        run = p.add_run(f"\t{desc}")
        set_run(run, size=10.5)

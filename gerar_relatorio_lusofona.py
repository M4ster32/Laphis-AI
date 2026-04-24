"""
Gera o relatório LAPHIS seguindo o template oficial da Universidade Lusófona
para Trabalho Final de Curso — Projeto II (Entrega Intercalar, Abril 2026).

Estrutura do template:
  Capa · Direitos de cópia · Agradecimentos · Resumo · Abstract
  Índice · Lista de Figuras · Lista de Tabelas · Lista de Siglas
  1. Introdução
  2. Pertinência e Viabilidade
  3. Especificação e Modelação
  4. Solução Desenvolvida
  5. Testes e Validação
  6. Método e Planeamento
  7. Resultados
  8. Conclusão
  Bibliografia
  Anexo 1
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "docs/LAPHIS_Relatorio_ProjetoII_Lusofona.docx"

# Paleta (tons sóbrios compatíveis com o template Lusófona)
NAVY = RGBColor(0x1F, 0x3A, 0x68)      # azul escuro do logo
TEXT = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x5A, 0x63)
LIGHT = RGBColor(0x8A, 0x90, 0x9A)

# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------


def set_run(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", size=11, bold=False, italic=False, color=None,
             align=None, space_after=6, space_before=0, first_line=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(first_line)
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1, number=None):
    """Heading com numeração manual (1, 1.1, 1.1.1)."""
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if number:
        run_num = p.add_run(f"{number}\t")
        set_run(run_num, size=sizes.get(level, 12), bold=True, color=NAVY)
    run = p.add_run(text)
    set_run(run, size=sizes.get(level, 12), bold=True, color=NAVY)
    # page break before para heading level 1
    if level == 1:
        p.paragraph_format.page_break_before = True
    return p


def add_bullet(doc, text, level=0, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.75)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, size=size)
    return p


def add_numbered(doc, text, size=11):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, size=size)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_caption(doc, text, kind="Figura"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"{kind} – {text}")
    set_run(run, size=9.5, italic=True, color=GREY)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # header shading (navy)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3A68")
        tcPr.append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            set_run(run, size=9.5)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    # spacing after
    doc.add_paragraph()
    return table


def add_code_block(doc, text, size=9):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F3F5")
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        set_run(run, name="Consolas", size=size, color=TEXT)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# DOCUMENT SETUP
# ---------------------------------------------------------------------------

doc = Document()

# Margens & estilos
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Header (título do trabalho em itálico, como no template)
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("LAPHIS — Life and Physical Health Intelligent System")
set_run(hr, size=10, italic=True, color=GREY)

# Footer com número de página
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fr = OxmlElement("w:r")
frPr = OxmlElement("w:rPr")
fsz = OxmlElement("w:sz"); fsz.set(qn("w:val"), "20")
frPr.append(fsz)
fr.append(frPr)
ft = OxmlElement("w:t"); ft.text = "1"
fr.append(ft)
fld.append(fr)
fp._p.append(fld)

# ===========================================================================
# CAPA
# ===========================================================================

for _ in range(3):
    doc.add_paragraph()

# "Universidade Lusófona" (placeholder — logo colocar-se-á manualmente)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UNIVERSIDADE")
set_run(r, name="Times New Roman", size=16, color=NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LUSÓFONA")
set_run(r, name="Times New Roman", size=28, bold=True, color=NAVY)

for _ in range(3):
    doc.add_paragraph()

# Título
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPHIS")
set_run(r, name="Times New Roman", size=36, bold=True, color=TEXT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Plataforma Inteligente de Saúde e Fitness com IA Conversacional")
set_run(r, name="Times New Roman", size=16, color=TEXT)

for _ in range(2):
    doc.add_paragraph()

# Bloco TFC
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Trabalho Final de Curso")
set_run(r, name="Times New Roman", size=14, bold=True, color=TEXT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Licenciatura em Engenharia Informática")
set_run(r, size=11, bold=True, color=TEXT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Projeto II")
set_run(r, size=11, bold=True, color=TEXT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Relatório Intercalar")
set_run(r, size=11, color=TEXT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2º Semestre")
set_run(r, size=11, color=TEXT)

for _ in range(2):
    doc.add_paragraph()

# Autor
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Henrique Fernandes")
set_run(r, size=11, bold=True, color=TEXT)

doc.add_paragraph()

# Orientador
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Orientador: ")
set_run(r, size=11, bold=True, color=TEXT)
r = p.add_run("[nome do orientador]")
set_run(r, size=11, color=TEXT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Entidade Externa: ")
set_run(r, size=11, bold=True, color=TEXT)
r = p.add_run("<se aplicável>")
set_run(r, size=11, color=TEXT)

for _ in range(2):
    doc.add_paragraph()

# Departamento
for line in [
    "Departamento de Engenharia Informática e Sistemas de Informação",
    "Universidade Lusófona, Centro Universitário do Porto (CUP)",
    "Abril 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    set_run(r, size=10.5, color=TEXT)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("w w w . u l u s o f o n a . p t")
set_run(r, size=9, color=NAVY)

page_break(doc)

# ===========================================================================
# DIREITOS DE CÓPIA
# ===========================================================================

for _ in range(10):
    doc.add_paragraph()

add_para(doc, "Direitos de cópia", size=14, bold=True, color=NAVY,
         space_after=12)
add_para(doc,
    "LAPHIS — Plataforma Inteligente de Saúde e Fitness com IA Conversacional, "
    "Copyright de Henrique Fernandes, Universidade Lusófona.",
    size=11, space_after=10)
add_para(doc,
    "A Faculdade de Ciências Naturais, Engenharia e Tecnologias (FCNET) e a "
    "Universidade Lusófona têm o direito, perpétuo e sem limites geográficos, "
    "de arquivar e publicar esta dissertação/monografia através de exemplares "
    "impressos reproduzidos em papel ou de forma digital, ou por qualquer outro "
    "meio conhecido ou que venha a ser inventado, e de a divulgar através de "
    "repositórios científicos e de admitir a sua cópia e distribuição com "
    "objetivos educacionais ou de investigação, não comerciais, desde que seja "
    "dado crédito ao autor e editor.",
    size=11)

page_break(doc)

# ===========================================================================
# AGRADECIMENTOS
# ===========================================================================

add_para(doc, "Agradecimentos", size=16, bold=True, color=NAVY, space_after=18)
add_para(doc,
    "Aos professores do Departamento de Engenharia Informática e Sistemas de "
    "Informação da Universidade Lusófona, pelo acompanhamento ao longo das "
    "unidades curriculares de Projeto I e Projeto II.",
    size=11, space_after=8)
add_para(doc,
    "Ao orientador deste trabalho, pela disponibilidade, sugestões e apreciação "
    "crítica nas várias fases de desenvolvimento do projeto.",
    size=11, space_after=8)
add_para(doc,
    "Aos colegas que testaram versões preliminares da aplicação e cujo "
    "feedback contribuiu para decisões de desenho e funcionalidade.",
    size=11, space_after=8)
add_para(doc,
    "À família e amigos, pelo apoio durante o percurso académico que culmina "
    "neste relatório.",
    size=11)

page_break(doc)

# ===========================================================================
# RESUMO
# ===========================================================================

add_para(doc, "Resumo", size=16, bold=True, color=NAVY, space_after=12)
add_para(doc,
    "O LAPHIS (Life and Physical Health Intelligent System) é uma plataforma web "
    "de gestão integrada de saúde e fitness, desenvolvida no âmbito do Trabalho "
    "Final de Curso da Licenciatura em Engenharia Informática. O problema em "
    "estudo é a fragmentação das ferramentas atuais de tracking pessoal — treino, "
    "nutrição, hidratação, peso e bem-estar mental encontram-se dispersos por "
    "aplicações distintas, sem cruzamento de dados nem recomendações adaptadas "
    "ao perfil individual. Para resolver este problema, foi desenvolvida uma "
    "aplicação cliente-servidor com frontend em React 19 (Vite) e backend em "
    "FastAPI (Python 3.11), assente em base de dados relacional gerida por "
    "SQLAlchemy 2.0. A inteligência do sistema é suportada por um motor de "
    "recomendação próprio de 857 linhas, baseado em heurísticas nutricionais e "
    "de exercício (BMI, TDEE, distribuição de macronutrientes, plano diário "
    "adaptativo), complementado por uma camada opcional de Inteligência "
    "Artificial (OpenAI GPT) com pipeline de Retrieval-Augmented Generation "
    "sobre documentos PDF. O sistema foi validado por uma suite automatizada de "
    "62 testes (pytest + httpx) e encontra-se atualmente em fase de deploy em "
    "produção (Vercel para o frontend, Render para o backend e base de dados "
    "PostgreSQL). O relatório documenta a arquitetura, tecnologias, "
    "especificação de requisitos, decisões de desenho, progresso até à entrega "
    "intercalar e trabalhos futuros previstos para a entrega final.",
    size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=14)

p = doc.add_paragraph()
r = p.add_run("Palavras-chave: ")
set_run(r, size=11, bold=True)
r = p.add_run("saúde digital; inteligência artificial; recomendação personalizada; "
              "aplicação web; FastAPI")
set_run(r, size=11)

page_break(doc)

# ===========================================================================
# ABSTRACT
# ===========================================================================

add_para(doc, "Abstract", size=16, bold=True, color=NAVY, space_after=12)
add_para(doc,
    "LAPHIS (Life and Physical Health Intelligent System) is an integrated web "
    "platform for health and fitness management, developed within the scope of "
    "the Final Course Project for the Bachelor's degree in Computer Engineering. "
    "The problem addressed is the fragmentation of current personal tracking "
    "tools — workouts, nutrition, hydration, weight and mental well-being are "
    "scattered across separate applications, without cross-domain data "
    "correlation or truly personalized recommendations. To tackle this, a "
    "client-server application was developed with a React 19 (Vite) frontend "
    "and a FastAPI (Python 3.11) backend, built on a relational database "
    "managed through SQLAlchemy 2.0. Intelligence is provided by a custom "
    "857-line recommendation engine based on nutritional and exercise "
    "heuristics (BMI, TDEE, macronutrient distribution, adaptive daily plan), "
    "complemented by an optional Artificial Intelligence layer (OpenAI GPT) "
    "with a Retrieval-Augmented Generation pipeline over PDF documents. The "
    "system is validated by an automated suite of 62 tests (pytest + httpx) "
    "and is currently being deployed to production (Vercel for the frontend, "
    "Render for the backend and PostgreSQL database). This report documents "
    "the architecture, technology stack, requirements specification, design "
    "decisions, progress up to the intercalar delivery and planned future "
    "work for the final delivery.",
    size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=14)

p = doc.add_paragraph()
r = p.add_run("Keywords: ")
set_run(r, size=11, bold=True)
r = p.add_run("digital health; artificial intelligence; personalized recommendation; "
              "web application; FastAPI")
set_run(r, size=11)

page_break(doc)

# ===========================================================================
# ÍNDICE (placeholder — gerado pelo Word)
# ===========================================================================

add_para(doc, "Índice", size=16, bold=True, color=NAVY, space_after=12)
add_para(doc,
    "[Este índice deve ser atualizado automaticamente no Word usando "
    "Referências → Índice. Clicar com o botão direito → Atualizar campos.]",
    size=10, italic=True, color=GREY, space_after=12)

# Índice manual como placeholder
toc_entries = [
    ("Agradecimentos", "iii"),
    ("Resumo", "iv"),
    ("Abstract", "v"),
    ("Índice", "vi"),
    ("Lista de Figuras", "viii"),
    ("Lista de Tabelas", "ix"),
    ("Lista de Siglas", "x"),
    ("1  Introdução", "1"),
    ("    1.1  Enquadramento", "1"),
    ("    1.2  Motivação e Identificação do Problema", "2"),
    ("    1.3  Objetivos", "3"),
    ("    1.4  Estrutura do Documento", "3"),
    ("2  Pertinência e Viabilidade", "4"),
    ("    2.1  Pertinência", "4"),
    ("    2.2  Viabilidade", "5"),
    ("    2.3  Análise Comparativa com Soluções Existentes", "5"),
    ("    2.4  Proposta de Inovação e Mais-Valias", "7"),
    ("    2.5  Identificação de Oportunidade de Negócio", "7"),
    ("3  Especificação e Modelação", "8"),
    ("    3.1  Análise de Requisitos", "8"),
    ("    3.2  Modelação", "11"),
    ("    3.3  Protótipos de Interface", "13"),
    ("4  Solução Desenvolvida", "14"),
    ("    4.1  Introdução", "14"),
    ("    4.2  Arquitetura", "14"),
    ("    4.3  Tecnologias e Ferramentas Utilizadas", "16"),
    ("    4.4  Ambientes de Teste e de Produção", "17"),
    ("    4.5  Abrangência", "18"),
    ("    4.6  Componentes", "18"),
    ("    4.7  Interfaces", "20"),
    ("5  Testes e Validação", "22"),
    ("6  Método e Planeamento", "24"),
    ("7  Resultados", "26"),
    ("    7.1  Resultados dos Testes", "26"),
    ("    7.2  Cumprimento de Requisitos", "27"),
    ("8  Conclusão", "28"),
    ("    8.1  Conclusão", "28"),
    ("    8.2  Trabalhos Futuros", "29"),
    ("Bibliografia", "30"),
    ("Anexo 1 — Recomendações para formatação de um relatório", "31"),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    run = p.add_run(f"{entry}\t{page}")
    set_run(run, size=10.5, color=TEXT)

page_break(doc)

# ===========================================================================
# LISTA DE FIGURAS
# ===========================================================================

add_para(doc, "Lista de Figuras", size=16, bold=True, color=NAVY, space_after=12)
figures = [
    ("Figura 1 – Arquitetura geral do sistema LAPHIS.", "15"),
    ("Figura 2 – Diagrama Entidade-Relação da base de dados.", "11"),
    ("Figura 3 – Fluxo de autenticação e verificação de email.", "19"),
    ("Figura 4 – Pipeline de RAG sobre documentos PDF.", "20"),
    ("Figura 5 – Mapa aplicacional e navegação entre ecrãs.", "13"),
    ("Figura 6 – Diagrama de casos de uso (ator utilizador).", "10"),
    ("Figura 7 – Cronograma Gantt do projeto.", "25"),
]
for entry, page in figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    run = p.add_run(f"{entry}\t{page}")
    set_run(run, size=10.5)

page_break(doc)

# ===========================================================================
# LISTA DE TABELAS
# ===========================================================================

add_para(doc, "Lista de Tabelas", size=16, bold=True, color=NAVY, space_after=12)
tables_list = [
    ("Tabela 1 – Análise de benchmarking comparativo com soluções existentes.", "6"),
    ("Tabela 2 – Requisitos funcionais identificados.", "8"),
    ("Tabela 3 – Requisitos não-funcionais identificados.", "10"),
    ("Tabela 4 – Entidades do modelo de dados.", "12"),
    ("Tabela 5 – Tecnologias utilizadas no backend.", "16"),
    ("Tabela 6 – Tecnologias utilizadas no frontend.", "16"),
    ("Tabela 7 – Plano de testes por módulo.", "22"),
    ("Tabela 8 – Cumprimento de requisitos.", "27"),
]
for entry, page in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    run = p.add_run(f"{entry}\t{page}")
    set_run(run, size=10.5)

page_break(doc)

# ===========================================================================
# LISTA DE SIGLAS
# ===========================================================================

add_para(doc, "Lista de Siglas", size=16, bold=True, color=NAVY, space_after=12)
siglas = [
    ("API", "Application Programming Interface"),
    ("BMI", "Body Mass Index (Índice de Massa Corporal)"),
    ("CI/CD", "Continuous Integration / Continuous Deployment"),
    ("CORS", "Cross-Origin Resource Sharing"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("CSS", "Cascading Style Sheets"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("IA", "Inteligência Artificial"),
    ("JSON", "JavaScript Object Notation"),
    ("JWT", "JSON Web Token"),
    ("LLM", "Large Language Model"),
    ("ORM", "Object-Relational Mapping"),
    ("PDF", "Portable Document Format"),
    ("RAG", "Retrieval-Augmented Generation"),
    ("REST", "Representational State Transfer"),
    ("SPA", "Single-Page Application"),
    ("SQL", "Structured Query Language"),
    ("TDEE", "Total Daily Energy Expenditure"),
    ("TFC", "Trabalho Final de Curso"),
    ("UML", "Unified Modeling Language"),
]
for sigla, desc in siglas:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(3.5))
    run = p.add_run(sigla)
    set_run(run, size=10.5, bold=True)
    run = p.add_run(f"\t{desc}")
    set_run(run, size=10.5)

# ===========================================================================
# 1. INTRODUÇÃO
# ===========================================================================

add_heading(doc, "Introdução", level=1, number="1")

add_heading(doc, "Enquadramento", level=2, number="1.1")
add_para(doc,
    "A adoção massiva de smartphones, dispositivos vestíveis (wearables) e "
    "aplicações móveis de autocuidado democratizou o acesso ao acompanhamento "
    "pessoal de indicadores de saúde. O utilizador médio regista, hoje, "
    "atividade física, ingestão calórica, padrões de sono e sinais biométricos "
    "através de uma combinação de ferramentas digitais. Contudo, o ecossistema "
    "que emergiu é profundamente fragmentado: cada aplicação domina um silo "
    "funcional (treino, nutrição, hidratação, mindfulness) e os dados raramente "
    "cruzam fronteiras, resultando em recomendações genéricas que ignoram o "
    "contexto integral do utilizador.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "Paralelamente, a evolução recente dos Modelos de Linguagem de Grande Escala "
    "(LLM) e das técnicas de Retrieval-Augmented Generation (RAG) abriu espaço "
    "para que sistemas personalizados cruzem dados biométricos, histórico "
    "individual e literatura científica, oferecendo aconselhamento adaptativo "
    "com um nível de granularidade que até recentemente não estava ao alcance "
    "de soluções de mercado [OpAI24, CiAl24].",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "É neste enquadramento que se insere o LAPHIS (Life and Physical Health "
    "Intelligent System): uma plataforma web integrada que agrega tracking "
    "multi-domínio, um motor de recomendação próprio baseado em heurísticas "
    "clínicas reconhecidas (equação de Harris-Benedict para cálculo de TDEE, "
    "fórmula BMI, distribuição de macronutrientes), e uma camada opcional de "
    "IA conversacional para aconselhamento personalizado.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Motivação e Identificação do Problema", level=2, number="1.2")
add_para(doc,
    "A motivação para o desenvolvimento do LAPHIS parte da observação — "
    "pessoal e documentada em estudos de utilização — de que o utilizador "
    "típico interage diariamente com entre três e cinco aplicações distintas "
    "para gerir a sua saúde e fitness. O custo cognitivo desta dispersão, "
    "associado à ausência de uma visão holística do perfil, traduz-se em "
    "abandono de hábitos e em resultados aquém do potencial.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "Três problemas concretos foram identificados como guia do trabalho:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc,
    "Fragmentação de dados — treino, nutrição e outros registos são mantidos em "
    "silos, impedindo o cruzamento necessário para recomendações integradas.")
add_bullet(doc,
    "Recomendações estáticas e genéricas — planos tabelados que ignoram "
    "objetivos individuais, restrições alimentares, nível real de atividade e "
    "feedback contínuo.")
add_bullet(doc,
    "Ausência de adaptação — planos que não reagem à evolução real do "
    "utilizador, à estagnação, ao progresso, ou à mudança de prioridades.")

add_heading(doc, "Objetivos", level=2, number="1.3")
add_para(doc,
    "O objetivo geral do trabalho é a conceção, implementação e disponibilização "
    "em ambiente produtivo de uma plataforma web completa de saúde e fitness com "
    "IA, que resolva os problemas de fragmentação, personalização e adaptação "
    "identificados. Os objetivos específicos dividem-se nos seguintes pontos:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "Desenvolver uma aplicação web cliente-servidor com separação "
                "clara entre SPA React e API REST FastAPI.")
add_bullet(doc, "Implementar um motor de recomendação próprio, suportado por "
                "heurísticas clínicas reconhecidas.")
add_bullet(doc, "Integrar uma camada opcional de IA (OpenAI) com pipeline RAG "
                "sobre PDFs, com fallback para modo baseado em regras.")
add_bullet(doc, "Garantir autenticação segura com verificação de email, hashing "
                "bcrypt e JWT stateless.")
add_bullet(doc, "Assegurar qualidade com suite de testes automatizados e código "
                "livre de dependências deprecated.")
add_bullet(doc, "Colocar o sistema em produção através de pipeline CI/CD "
                "automatizado (Vercel + Render).")

add_heading(doc, "Estrutura do Documento", level=2, number="1.4")
add_para(doc,
    "O presente relatório segue a estrutura recomendada para o Trabalho Final "
    "de Curso e está organizado da seguinte forma:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "Na Secção 2 é discutida a pertinência e viabilidade do trabalho, "
                "incluindo análise comparativa com soluções existentes.")
add_bullet(doc, "Na Secção 3 apresentam-se a especificação e modelação — requisitos "
                "funcionais e não-funcionais, casos de uso, diagrama entidade-relação "
                "e protótipos de interface.")
add_bullet(doc, "Na Secção 4 descreve-se a solução desenvolvida: arquitetura, stack "
                "tecnológico, componentes e interfaces.")
add_bullet(doc, "Na Secção 5 detalha-se o plano de testes e validação.")
add_bullet(doc, "Na Secção 6 apresenta-se o método e planeamento, com cronograma.")
add_bullet(doc, "Na Secção 7 documentam-se os resultados preliminares dos testes e "
                "o grau de cumprimento dos requisitos.")
add_bullet(doc, "Na Secção 8 conclui-se o relatório, identificando os trabalhos "
                "futuros previstos para a entrega final.")
add_bullet(doc, "No Anexo 1 incluem-se elementos de formatação e referência.")

# ===========================================================================
# 2. PERTINÊNCIA E VIABILIDADE
# ===========================================================================

add_heading(doc, "Pertinência e Viabilidade", level=1, number="2")

add_heading(doc, "Pertinência", level=2, number="2.1")
add_para(doc,
    "O impacto positivo esperado do LAPHIS assenta em três pilares. Primeiro, "
    "redução do custo cognitivo de gestão da saúde pessoal: ao consolidar "
    "múltiplos domínios numa única aplicação, elimina-se a necessidade de "
    "alternar entre ferramentas e reconciliar manualmente dados dispersos. "
    "Segundo, personalização real: um motor de recomendação que cruza perfil "
    "biométrico, objetivos declarados e registos históricos gera aconselhamento "
    "adaptado, em contraste com os planos genéricos oferecidos pela maioria das "
    "soluções de mercado. Terceiro, acessibilidade tecnológica: o sistema é "
    "entregue como aplicação web, dispensando instalação e funcionando em "
    "qualquer dispositivo com browser moderno.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "A validação por terceiros desta pertinência foi iniciada de forma informal "
    "junto de colegas e utilizadores-teste durante o desenvolvimento. Nesta fase "
    "intercalar, o feedback recolhido será sistematizado e apresentado em "
    "inquérito estruturado na entrega final.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Viabilidade", level=2, number="2.2")
add_para(doc,
    "A viabilidade técnica do projeto está demonstrada pelo estado atual de "
    "implementação: a aplicação encontra-se funcional em ambiente de "
    "desenvolvimento, com todos os módulos principais (autenticação, tracking, "
    "planos, chat) operacionais, e o deploy em produção em execução à data "
    "desta entrega intercalar. A utilização de ferramentas gratuitas ou com "
    "tiers gratuitos generosos (Vercel, Render, GitHub, PostgreSQL free tier) "
    "garante a sustentabilidade financeira do projeto enquanto trabalho "
    "académico. A componente de IA é opcional: na ausência de chave OpenAI, o "
    "sistema funciona integralmente em modo baseado em regras.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "Do ponto de vista económico, um eventual cenário comercial apoia-se em "
    "modelo freemium — funcionalidades essenciais gratuitas, funcionalidades "
    "premium (assistente IA avançado, integração com wearables, histórico "
    "ilimitado) mediante subscrição. A manutenção e evolução após o término do "
    "TFC é viável, dado o caráter modular da arquitetura e a cobertura de "
    "testes existente.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Análise Comparativa com Soluções Existentes", level=2,
            number="2.3")

add_heading(doc, "Soluções existentes", level=3, number="2.3.1")
add_para(doc,
    "Foram identificadas no mercado quatro soluções principais, cada uma "
    "especializada num subconjunto do domínio de saúde e fitness:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc,
    "MyFitnessPal — tracking de nutrição com base de dados alimentar extensa; "
    "funcionalidades de treino e integração limitadas.")
add_bullet(doc,
    "Strava — rede social e tracking de atividade física outdoor; sem "
    "componente de nutrição ou bem-estar mental.")
add_bullet(doc,
    "Calm / Headspace — meditação guiada e mindfulness; sem integração com "
    "tracking físico ou nutricional.")
add_bullet(doc,
    "Fitbit / Apple Health — hubs de dados biométricos agregados por "
    "dispositivos vestíveis; fracos em recomendação e planeamento.")

add_heading(doc, "Análise de benchmarking", level=3, number="2.3.2")
add_para(doc,
    "A Tabela 1 sistematiza a comparação da solução LAPHIS com as alternativas "
    "identificadas, sinalizando com \"X\" a presença de cada característica-chave.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_caption(doc, "Análise de benchmarking comparativo com soluções existentes.",
            kind="Tabela 1")
add_table(doc,
    ["Característica", "MyFitnessPal", "Strava", "Calm", "Fitbit", "LAPHIS"],
    [
        ["Tracking treino", "", "X", "", "X", "X"],
        ["Tracking nutrição", "X", "", "", "parcial", "X"],
        ["Hidratação", "parcial", "", "", "parcial", "X"],
        ["Peso corporal", "X", "", "", "X", "X"],
        ["Meditação", "", "", "X", "parcial", "X"],
        ["Planos personalizados", "parcial", "parcial", "", "", "X"],
        ["Chat IA", "", "", "", "", "X"],
        ["RAG sobre PDFs", "", "", "", "", "X"],
        ["Dark mode nativo", "parcial", "X", "X", "X", "X"],
        ["Open source / académico", "", "", "", "", "X"],
    ],
    col_widths=[3.5, 2.5, 2.0, 2.0, 2.0, 2.0])

add_heading(doc, "Proposta de Inovação e Mais-Valias", level=2, number="2.4")
add_para(doc,
    "O elemento distintivo do LAPHIS é a integração cross-domínio combinada "
    "com uma camada de IA opcional e um motor de regras transparente. A "
    "grande maioria das soluções comerciais opera num único silo, deixando ao "
    "utilizador a responsabilidade de cruzar mentalmente os dados. O LAPHIS "
    "inverte este paradigma: os registos de treino influenciam as "
    "recomendações nutricionais, o peso corporal alimenta o cálculo de "
    "hidratação recomendada, e o feedback sobre planos informa ajustes "
    "futuros.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "Mais-valias adicionais: (i) transparência do motor de recomendação — as "
    "fórmulas usadas (Harris-Benedict, distribuição de macros) são publicadas "
    "no código, ao contrário do caráter opaco da maioria das aplicações; "
    "(ii) pipeline RAG sobre PDFs, que permite ao utilizador carregar "
    "documentos próprios (relatórios médicos, planos nutricionais de "
    "profissionais) e fazer consultas em linguagem natural; (iii) design "
    "system premium com dark/light mode completo; (iv) disponibilização em "
    "ambiente produtivo real.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Identificação de Oportunidade de Negócio", level=2,
            number="2.5")
add_para(doc,
    "Caso se pretenda transpor o LAPHIS para exploração comercial, o modelo de "
    "negócio natural é freemium com segmento B2C e extensão B2B a prazo:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "Tier gratuito — tracking completo, chat com IA limitado a N "
                "sessões/mês, planos gerados por regras.")
add_bullet(doc, "Tier Premium (€4,99/mês) — chat IA ilimitado, planos gerados "
                "por LLM, integração com wearables, histórico ilimitado.")
add_bullet(doc, "Tier Profissional — marketplace de planos validados por "
                "nutricionistas e personal trainers certificados, com "
                "comissão sobre transações.")
add_bullet(doc, "Licenciamento B2B — ginásios e clínicas podem oferecer "
                "LAPHIS white-label aos seus clientes.")

# ===========================================================================
# 3. ESPECIFICAÇÃO E MODELAÇÃO
# ===========================================================================

add_heading(doc, "Especificação e Modelação", level=1, number="3")

add_heading(doc, "Análise de Requisitos", level=2, number="3.1")
add_para(doc,
    "A análise de requisitos resulta do cruzamento da investigação sobre "
    "soluções existentes (Secção 2.3), das especificações da UC de Análise "
    "de Sistemas e de conversas iterativas com potenciais utilizadores. A "
    "enumeração inicial, estabelecida em Projeto I, foi revista nesta entrega "
    "intercalar — os requisitos modificados são indicados explicitamente.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Requisitos Funcionais", level=3, number="3.1.1")
add_para(doc,
    "A Tabela 2 apresenta a lista de requisitos funcionais identificados, com "
    "indicação do estado atual de implementação. O estado \"Realizado\" indica "
    "implementação completa e testada; \"Parcial\" indica funcionalidade básica "
    "existente mas com refinamentos pendentes para a entrega final.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_caption(doc, "Requisitos funcionais identificados.", kind="Tabela 2")
add_table(doc,
    ["ID", "Requisito", "Descrição", "Estado"],
    [
        ["RF01", "Registo e autenticação", "Utilizador cria conta com email/password, "
         "confirma por código enviado por email, autentica-se obtendo JWT.", "Realizado"],
        ["RF02", "Gestão de perfil", "Utilizador gere dados biométricos (idade, "
         "peso, altura, objetivo, nível de atividade, dieta).", "Realizado"],
        ["RF03", "Registo de treinos", "Utilizador regista treinos com exercício, "
         "duração, calorias, categoria.", "Realizado"],
        ["RF04", "Registo de refeições", "Utilizador regista refeições com "
         "alimento, calorias, macronutrientes e tipo.", "Realizado"],
        ["RF05", "Tracking hidratação", "Registo de ingestão de água em ml, com "
         "meta diária calculada.", "Realizado"],
        ["RF06", "Tracking de peso", "Histórico de peso corporal com cálculo "
         "automático de BMI e tendência.", "Realizado"],
        ["RF07", "Planos de treino e nutrição", "Geração e gestão de planos "
         "personalizados, com feedback.", "Realizado"],
        ["RF08", "Plano diário adaptativo", "Plano do dia gerado automaticamente "
         "com treinos e refeições, ajustável.", "Parcial"],
        ["RF09", "Chat IA conversacional", "Sessões de chat com assistente IA "
         "com contexto do perfil.", "Realizado"],
        ["RF10", "Relatórios e gráficos", "Visualização semanal de progresso e "
         "exportação PDF.", "Realizado"],
        ["RF11", "Sessões de meditação", "Registo de sessões Zen com humor "
         "antes/depois.", "Realizado"],
        ["RF12", "Categorias personalizáveis", "Utilizador define categorias "
         "próprias de treino e refeição.", "Realizado"],
        ["RF13", "Pipeline RAG sobre PDFs", "Upload de documentos PDF, extração "
         "de texto e consulta por linguagem natural.", "Parcial"],
        ["RF14", "Notificações push", "Lembretes de registo de refeições, água "
         "e treino.", "Planeado"],
        ["RF15", "Integração wearables", "Sincronização com Apple Health, "
         "Fitbit, Google Fit.", "Planeado"],
    ],
    col_widths=[1.2, 3.5, 7.5, 2.0])

add_heading(doc, "Requisitos Não-Funcionais", level=3, number="3.1.2")
add_para(doc,
    "A Tabela 3 apresenta os requisitos não-funcionais que governam a "
    "qualidade do sistema em aspetos transversais — desempenho, segurança, "
    "usabilidade e manutenibilidade.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_caption(doc, "Requisitos não-funcionais identificados.", kind="Tabela 3")
add_table(doc,
    ["ID", "Categoria", "Descrição", "Estado"],
    [
        ["RNF01", "Performance", "Resposta da API < 2 s em operações típicas.", "Realizado"],
        ["RNF02", "Segurança", "Passwords com hash bcrypt, JWT stateless, HTTPS "
         "em produção.", "Realizado"],
        ["RNF03", "Responsividade", "Interface adaptável a mobile (<640px), "
         "tablet e desktop.", "Realizado"],
        ["RNF04", "Disponibilidade", "Uptime esperado ≥ 99% em produção, com "
         "health checks automáticos.", "Realizado"],
        ["RNF05", "Testabilidade", "Cobertura mínima de 50 testes unitários e "
         "de integração.", "Realizado (62)"],
        ["RNF06", "Manutenibilidade", "Código sem deprecations, separação em "
         "camadas, linting ativo.", "Realizado"],
        ["RNF07", "Acessibilidade", "Ratios WCAG AA para contraste, modo escuro "
         "completo.", "Realizado"],
        ["RNF08", "Escalabilidade", "Arquitetura stateless, base de dados "
         "substituível (SQLite ↔ PostgreSQL).", "Realizado"],
        ["RNF09", "Internacionalização", "Interface em português, mensagens "
         "isoladas para futura i18n.", "Parcial"],
        ["RNF10", "Privacidade", "Dados pessoais apenas na BD; nenhum envio a "
         "terceiros sem consentimento.", "Realizado"],
    ],
    col_widths=[1.2, 2.5, 8.5, 2.0])

add_heading(doc, "Casos de Uso", level=3, number="3.1.3")
add_para(doc,
    "Os casos de uso foram modelados tendo como ator primário o "
    "\"Utilizador registado\" e, como atores secundários, o \"Visitante "
    "(não registado)\" e o \"Assistente IA\". A Figura 6 esquematiza os casos "
    "de uso principais.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "Os principais casos de uso identificados são:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "UC01 — Registar conta e verificar email")
add_bullet(doc, "UC02 — Autenticar e manter sessão")
add_bullet(doc, "UC03 — Preencher perfil biométrico")
add_bullet(doc, "UC04 — Registar treino")
add_bullet(doc, "UC05 — Registar refeição")
add_bullet(doc, "UC06 — Registar hidratação")
add_bullet(doc, "UC07 — Registar peso e visualizar tendência")
add_bullet(doc, "UC08 — Gerar plano personalizado")
add_bullet(doc, "UC09 — Conversar com assistente IA")
add_bullet(doc, "UC10 — Visualizar relatório semanal e exportar PDF")
add_bullet(doc, "UC11 — Carregar documento PDF e consultar conteúdo")
add_bullet(doc, "UC12 — Registar sessão de meditação")

add_para(doc,
    "[Espaço reservado para Figura 6 — Diagrama de casos de uso. A incluir "
    "diagrama UML elaborado em ferramenta externa (PlantUML/Draw.io).]",
    italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading(doc, "Modelação", level=2, number="3.2")
add_para(doc,
    "O modelo de dados segue o paradigma relacional, com o utilizador (User) "
    "como entidade raiz. Todas as demais entidades mantêm chave estrangeira "
    "para User, com cascade configurado para permitir eliminação da conta com "
    "limpeza automática dos registos associados. Campos de data/hora são "
    "timezone-aware; estruturas semi-estruturadas (metrics, plan_data) usam "
    "campos JSON para manter flexibilidade.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "A Figura 2 apresenta o diagrama entidade-relação completo. A Tabela 4 "
    "detalha cada entidade, campos relevantes e propósito.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc,
    "[Espaço reservado para Figura 2 — Diagrama Entidade-Relação. A substituir "
    "por diagrama ER exportado do dbdiagram.io ou DrawSQL.]",
    italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

add_caption(doc, "Entidades do modelo de dados.", kind="Tabela 4")
add_table(doc,
    ["Entidade", "Campos principais", "Propósito"],
    [
        ["User", "id, email, hashed_password, is_verified, verify_code, reset_token",
         "Autenticação e identidade raiz."],
        ["Profile", "user_id, age, weight, height, goal, activity, diet_type",
         "Dados biométricos do utilizador."],
        ["Plan", "user_id, title, type, content, is_active, feedback",
         "Planos personalizados gerados."],
        ["ChatSession", "user_id, title, created_at, expires_at",
         "Sessões de chat com IA."],
        ["ChatMessage", "session_id, role, content, created_at",
         "Mensagens trocadas no chat."],
        ["WorkoutLog", "user_id, exercise, duration, calories, category, logged_at",
         "Registos de treino."],
        ["MealLog", "user_id, food_name, calories, protein, carbs, fat, meal_type",
         "Registos de refeição."],
        ["WaterLog", "user_id, amount_ml, logged_at",
         "Registos de hidratação."],
        ["WeightEntry", "user_id, weight_kg, logged_at",
         "Histórico de peso corporal."],
        ["ZenSession", "user_id, type, duration_min, mood_before, mood_after, notes",
         "Sessões de meditação."],
        ["Category", "user_id, name, type, icon, color",
         "Categorias personalizáveis."],
        ["DailyPlan", "user_id, date, plan_data, status",
         "Planeamento diário adaptativo."],
        ["ProgressSnapshot", "user_id, date, metrics",
         "Snapshots periódicos de progresso."],
        ["AdaptationLog", "user_id, trigger, old_params, new_params, reason",
         "Log de adaptações automáticas."],
        ["PlanFeedback", "plan_id, rating, comment, created_at",
         "Feedback dos utilizadores sobre planos."],
        ["WeeklySummary", "user_id, week_start, avg_calories, total_workouts, recommendations",
         "Resumos semanais agregados."],
        ["RAGDocument", "user_id, filename, text_chunks, uploaded_at",
         "Documentos PDF indexados para RAG."],
    ],
    col_widths=[3.0, 7.0, 5.0])

add_heading(doc, "Protótipos de Interface", level=2, number="3.3")
add_para(doc,
    "O mapa aplicacional compreende 17 ecrãs organizados em duas grandes áreas: "
    "área pública (landing page, fluxos de autenticação) e área privada (sob "
    "autenticação JWT). A Figura 5 representa a navegação entre ecrãs.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "Os protótipos seguem um sistema de design denominado \"Premium "
    "Monochromatic\", inspirado em ferramentas como Linear, Notion e Stripe, "
    "com suporte completo a dark/light mode implementado via CSS Custom "
    "Properties. A tipografia escala-se em seis níveis (Display 36pt → Micro "
    "12pt) e o espaçamento obedece a um grid de 4 pontos.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc,
    "[Espaço reservado para Figura 5 — Mapa aplicacional com screenshots "
    "reduzidos dos ecrãs principais. A incluir protótipos Figma exportados.]",
    italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ===========================================================================
# 4. SOLUÇÃO DESENVOLVIDA
# ===========================================================================

add_heading(doc, "Solução Desenvolvida", level=1, number="4")

add_heading(doc, "Introdução", level=2, number="4.1")
add_para(doc,
    "A solução desenvolvida, LAPHIS, é uma aplicação web de arquitetura "
    "cliente-servidor, com frontend Single-Page Application (SPA) em React 19 "
    "e backend REST em FastAPI. A comunicação é feita via HTTP/JSON com "
    "autenticação JWT stateless; o backend comunica com uma base de dados "
    "relacional (SQLite em desenvolvimento, PostgreSQL em produção) através "
    "do ORM SQLAlchemy 2.0.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "Face às soluções apresentadas no benchmarking (Secção 2.3), o LAPHIS "
    "distingue-se pela integração de múltiplos domínios num único produto, pela "
    "existência de um motor de recomendação baseado em heurísticas clínicas "
    "transparentes, e pela camada de IA opcional com pipeline RAG.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

add_para(doc, "Artefactos da entrega final (a completar):",
         bold=True, space_after=4)
add_bullet(doc, "Link para vídeo demonstrativo: [a adicionar antes da entrega final].")
add_bullet(doc, "Repositório Git público: https://github.com/M4ster32/Laphis.git")
add_bullet(doc, "Solução funcional em produção: https://laphis.vercel.app "
                "(credenciais de demonstração a fornecer antes da entrega final).")

add_para(doc,
    "As próximas subsecções detalham a arquitetura, o stack tecnológico, os "
    "ambientes de execução, a abrangência curricular do trabalho, os "
    "componentes implementados e os principais ecrãs da interface.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Arquitetura", level=2, number="4.2")
add_para(doc,
    "A arquitetura do LAPHIS divide-se em quatro camadas principais, "
    "representadas na Figura 1. A separação entre camadas foi orientada por "
    "princípios de responsabilidade única, stateless por omissão (exceto na "
    "persistência) e independência tecnológica — o frontend e o backend "
    "comunicam exclusivamente por API REST, permitindo substituir qualquer "
    "das camadas sem afetar as restantes.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_code_block(doc,
    "┌──────────────┐   HTTPS   ┌──────────────┐   TCP    ┌──────────────┐\n"
    "│   Browser    │ ────────▶ │    Vercel    │ ───────▶ │    Render    │\n"
    "│  React 19    │  JSON/JWT │  Edge CDN    │          │ FastAPI app  │\n"
    "└──────────────┘           └──────────────┘          └──────┬───────┘\n"
    "                                                            │\n"
    "                                                     ┌──────▼───────┐\n"
    "                                                     │  PostgreSQL  │\n"
    "                                                     │  (Render)    │\n"
    "                                                     └──────────────┘")

add_caption(doc, "Arquitetura geral do sistema LAPHIS.", kind="Figura 1")

add_para(doc,
    "As opções arquiteturais foram fundamentadas nos seguintes pontos:",
    space_after=4)
add_bullet(doc, "SPA em vez de server-rendered pages: permite interações ricas "
                "e fluídas, reduz roundtrips e simplifica o deploy separado "
                "(frontend estático em CDN).")
add_bullet(doc, "REST em vez de GraphQL: a complexidade do domínio é moderada "
                "e REST tem menor overhead de ferramentas para um projeto "
                "académico.")
add_bullet(doc, "JWT stateless em vez de sessões em cookies server-side: "
                "facilita o escalamento horizontal e remove a necessidade "
                "de store de sessões.")
add_bullet(doc, "SQLAlchemy com SQLite em dev e PostgreSQL em prod: o ORM "
                "abstrai a diferença, permitindo ciclo rápido de desenvolvimento.")

add_heading(doc, "Tecnologias e Ferramentas Utilizadas", level=2, number="4.3")
add_para(doc,
    "A Tabela 5 e a Tabela 6 listam as tecnologias principais de backend e "
    "frontend, respetivamente, com justificação do uso.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_caption(doc, "Tecnologias utilizadas no backend.", kind="Tabela 5")
add_table(doc,
    ["Tecnologia", "Versão", "Justificação"],
    [
        ["Python", "3.11+", "Linguagem madura com ecossistema rico em IA/ML."],
        ["FastAPI", "0.104+", "Framework async moderno, tipado, com OpenAPI auto."],
        ["SQLAlchemy", "2.0", "ORM de referência em Python, sintaxe 2.0 type-safe."],
        ["Pydantic", "2.x", "Validação de dados com type hints, serialização JSON."],
        ["Uvicorn", "0.24+", "Servidor ASGI de produção, compatível com FastAPI."],
        ["PyJWT", "2.12+", "Geração e validação de tokens JWT."],
        ["bcrypt", "4.1+", "Hashing seguro de passwords com salt."],
        ["OpenAI SDK", "1.14+", "Cliente oficial para API GPT."],
        ["PyMuPDF", "1.24+", "Extração de texto de PDFs para pipeline RAG."],
        ["pytest + httpx", "—", "Framework de testes com cliente HTTP async."],
    ],
    col_widths=[3.5, 2.0, 10.5])

add_caption(doc, "Tecnologias utilizadas no frontend.", kind="Tabela 6")
add_table(doc,
    ["Tecnologia", "Versão", "Justificação"],
    [
        ["React", "19.2", "Concurrent rendering, ecossistema dominante."],
        ["Vite", "8.0 (beta)", "Build tool com HMR instantâneo via ESBuild."],
        ["React Router", "7.13", "Routing declarativo com layouts aninhados."],
        ["Lucide React", "0.577", "300+ ícones SVG consistentes."],
        ["Recharts", "3.8", "Componentes de gráficos SVG responsivos."],
        ["jsPDF", "4.2", "Geração de relatórios PDF no cliente."],
        ["ESLint", "—", "Linting JavaScript/React com regras modernas."],
    ],
    col_widths=[3.5, 2.0, 10.5])

add_heading(doc, "Ambientes de Teste e de Produção", level=2, number="4.4")
add_para(doc,
    "O projeto utiliza três ambientes distintos, com configuração controlada "
    "por variáveis de ambiente:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "Desenvolvimento local — Python virtualenv, base de dados "
                "SQLite em ficheiro local, frontend servido pelo Vite dev "
                "server em http://localhost:5173.")
add_bullet(doc, "Testes automatizados — base de dados SQLite in-memory, "
                "cliente HTTP httpx assíncrono, fixtures pytest com "
                "isolamento total por teste.")
add_bullet(doc, "Produção — frontend deployado na Vercel com CDN global, "
                "backend na Render com 1 vCPU / 512 MB RAM, base de dados "
                "PostgreSQL gerida (tier gratuito Render, 1 GB de "
                "armazenamento).")

add_para(doc,
    "Os recursos atuais são suficientes para a fase de demonstração e "
    "validação por terceiros. Num cenário de exploração comercial, "
    "recomenda-se escalar o backend para tier pago (2 vCPU, 2 GB RAM) e "
    "migrar a base de dados para plano Standard (10 GB, backups diários).",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Abrangência", level=2, number="4.5")
add_para(doc,
    "O desenvolvimento do LAPHIS aplica e integra conteúdos das seguintes "
    "unidades curriculares da Licenciatura em Engenharia Informática:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "Programação Web — React, CSS, responsividade, consumo de APIs.")
add_bullet(doc, "Bases de Dados — modelação relacional, normalização, SQL.")
add_bullet(doc, "Engenharia de Software — requisitos, UML, testes, ciclo de vida.")
add_bullet(doc, "Análise de Sistemas — levantamento de requisitos, casos de uso, "
                "benchmarking.")
add_bullet(doc, "Sistemas Distribuídos — arquitetura cliente-servidor, REST, "
                "autenticação stateless.")
add_bullet(doc, "Segurança Informática — hashing bcrypt, JWT, HTTPS, CORS.")
add_bullet(doc, "Inteligência Artificial — integração de LLM, pipeline RAG, "
                "heurísticas de recomendação.")

add_heading(doc, "Componentes", level=2, number="4.6")

add_heading(doc, "Backend — API REST com FastAPI", level=3, number="4.6.1")
add_para(doc,
    "O backend organiza-se em 19 routers FastAPI, cada um responsável por "
    "um domínio funcional. O ponto de entrada é o ficheiro main.py, que "
    "regista todos os routers, configura CORS, e utiliza um lifespan handler "
    "para inicialização da base de dados e migrações leves.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "Os routers implementados são: auth (14 endpoints), plans (11), "
    "adaptation (11), chat (8), progress (8), reports (7), logs (6), "
    "exercises (6), daily_plan (6), zen (5), weight (5), water (5), "
    "categories (5), rag_ingest (4), rag_ask (4), profile (4), ask (3), "
    "ingest (1) e health (1), totalizando mais de 100 endpoints.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "O motor de recomendação (core/recommender.py, 857 linhas) implementa as "
    "heurísticas clínicas reconhecidas: BMI (peso/altura²), TDEE "
    "(equação de Harris-Benedict revista multiplicada pelo fator de "
    "atividade), distribuição de macronutrientes por objetivo, hidratação "
    "recomendada (30-35 ml/kg) e planeamento diário.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "A camada de IA (core/ai_engine.py e core/rag.py) integra a API OpenAI. "
    "A deteção da variável de ambiente OPENAI_API_KEY é feita no arranque: "
    "se presente, o modo IA é ativado; caso contrário, o sistema opera em "
    "modo baseado em regras sem qualquer falha. O pipeline RAG extrai texto "
    "de PDFs via PyMuPDF, segmenta em chunks e injeta o contexto relevante "
    "no prompt enviado ao LLM.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Frontend — SPA React", level=3, number="4.6.2")
add_para(doc,
    "O frontend organiza-se em 17 páginas e 13 componentes reutilizáveis. O "
    "routing, gerido pelo React Router 7, distingue rotas públicas (landing, "
    "autenticação) de rotas protegidas (toda a área /app), estas últimas "
    "exigindo token JWT válido no contexto.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "A gestão de estado é feita via React Context API em dois contextos "
    "globais: AppContext (token, perfil do utilizador, flags de loading) e "
    "ThemeContext (modo dark/light, persistência em localStorage). A "
    "comunicação HTTP é centralizada em services/api.js, um wrapper que "
    "injeta automaticamente o header Authorization: Bearer <token> em todos "
    "os pedidos autenticados.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "Os componentes reutilizáveis incluem: Button, Card, Modal, Form, Toast, "
    "Skeleton, Icon, AvatarPicker, PasswordInput, EmptyState, ExerciseCard, "
    "GeneratePlanModal e um index.js que consolida os exports.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Interfaces", level=2, number="4.7")
add_para(doc,
    "Os ecrãs mais representativos da aplicação — a incluir aqui como "
    "screenshots na entrega final — são:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "Dashboard — visão geral do dia com KPIs de calorias, macros, "
                "água e peso; plano do dia abaixo.")
add_bullet(doc, "Logs — registo unificado de treinos e refeições, com filtros "
                "por data e categoria.")
add_bullet(doc, "Plans / PlanDetail — lista de planos ativos e detalhe "
                "editável com geração por IA ou regras.")
add_bullet(doc, "Chat — interface conversacional com bubbles, histórico por "
                "sessão e input com suporte a multi-linha.")
add_bullet(doc, "Reports — gráficos Recharts (linhas, barras, radial) e botão "
                "de exportação PDF.")
add_bullet(doc, "Zen — seleção de tipo de meditação, timer circular e "
                "registo de humor antes/depois.")

add_para(doc,
    "[Espaço reservado para screenshots dos ecrãs principais. A incluir na "
    "entrega final com capturas em modo escuro em resolução 1440×900.]",
    italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ===========================================================================
# 5. TESTES E VALIDAÇÃO
# ===========================================================================

add_heading(doc, "Testes e Validação", level=1, number="5")

add_para(doc,
    "O plano de testes do LAPHIS é construído em dois níveis complementares: "
    "testes automatizados do backend, focados em correção funcional e regressão, "
    "e testes de aceitação por utilizadores, orientados à validação operacional "
    "em contexto real. O plano de testes foi desenhado em paralelo com a "
    "especificação dos requisitos (Secção 3.1), de modo a garantir cobertura "
    "dos comportamentos críticos do sistema.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc, "Testes automatizados — cobertura atual", bold=True, space_after=4)
add_para(doc,
    "A suite atual reúne 62 testes automatizados, executados via pytest com "
    "cliente HTTP assíncrono httpx. Cada teste corre contra uma base de dados "
    "SQLite in-memory criada e destruída por teste, garantindo isolamento total.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_caption(doc, "Plano de testes por módulo.", kind="Tabela 7")
add_table(doc,
    ["Ficheiro", "Testes", "Cobertura funcional"],
    [
        ["test_auth.py", "~10", "Registo, login, verificação email, reset password, "
         "tokens inválidos, conta não verificada."],
        ["test_profile.py", "~8", "Criar, atualizar, obter perfil; validação de "
         "campos biométricos; cálculos automáticos."],
        ["test_logs.py", "~12", "CRUD de treinos e refeições; listagem paginada; "
         "filtros por data e categoria."],
        ["test_water.py", "~8", "Registo de água, total diário, histórico, meta "
         "baseada no peso."],
        ["test_weight.py", "~8", "Registo de peso, histórico, tendência, cálculo "
         "automático de BMI."],
        ["test_plans.py", "~10", "Gerar, guardar, listar, atualizar, eliminar "
         "planos; envio de feedback."],
        ["test_zen.py", "~6", "Criar sessão, listar, estatísticas, tipos de "
         "meditação."],
    ],
    col_widths=[3.5, 2.0, 10.5])

add_para(doc, "Testes de aceitação por terceiros (planeados)",
         bold=True, space_after=4)
add_para(doc,
    "Para a entrega final será aplicado um inquérito de satisfação a 10-15 "
    "utilizadores-teste, combinando respostas Likert e perguntas abertas, com "
    "o guião detalhado no Anexo de testes. Os critérios de avaliação incluem:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "Usabilidade — facilidade de registo, navegação, realização das "
                "tarefas principais.")
add_bullet(doc, "Percepção de utilidade — os utilizadores sentem que o LAPHIS "
                "resolve o problema de fragmentação identificado?")
add_bullet(doc, "Qualidade das recomendações — os planos e sugestões do "
                "assistente IA são percebidos como personalizados?")
add_bullet(doc, "Desempenho percebido — tempo de resposta aceitável?")
add_bullet(doc, "Disponibilidade de funcionalidades — falta alguma funcionalidade "
                "considerada essencial?")

add_para(doc,
    "A análise de risco é realizada com recurso a diagrama de causa-efeito "
    "(Ishikawa), identificando os principais modos de falha: (i) "
    "indisponibilidade do backend (mitigação: health checks e auto-restart), "
    "(ii) expiração de token JWT em operação longa (mitigação: refresh "
    "automático no cliente), (iii) falha na API OpenAI (mitigação: fallback "
    "para modo regras), (iv) limites de tier gratuito da base de dados "
    "(mitigação: monitorização de tamanho e limpeza de sessões expiradas).",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ===========================================================================
# 6. MÉTODO E PLANEAMENTO
# ===========================================================================

add_heading(doc, "Método e Planeamento", level=1, number="6")

add_para(doc,
    "A metodologia adotada é híbrida, combinando elementos de metodologias "
    "ágeis (desenvolvimento iterativo, entregas frequentes de incrementos "
    "funcionais) com fases clássicas em cascata para a documentação — "
    "levantamento de requisitos, modelação e testes finais. O trabalho "
    "individual exclui a necessidade de cerimónias formais de equipa, mas as "
    "práticas de controlo de versões (branches, commits atómicos, mensagens "
    "convencionais) foram mantidas com rigor.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc,
    "O desenvolvimento foi dividido em dois semestres letivos, correspondentes "
    "às UCs Projeto I (1º semestre, 2025/2026) e Projeto II (2º semestre, "
    "2025/2026). Os marcos principais são:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

add_para(doc, "Projeto I — 1º Semestre (Set 2025 – Jan 2026)",
         bold=True, space_after=4)
add_bullet(doc, "Set–Out 2025: Levantamento do estado da arte, estudo de "
                "soluções concorrentes, definição do problema.")
add_bullet(doc, "Nov 2025: Entrega intercalar — especificação de requisitos e "
                "modelação preliminar.")
add_bullet(doc, "Dez 2025 – Jan 2026: Implementação do núcleo funcional "
                "(autenticação, tracking básico, primeiros endpoints).")
add_bullet(doc, "Jan 2026: Entrega final de Projeto I.")

add_para(doc, "Projeto II — 2º Semestre (Fev 2026 – Mai 2026)",
         bold=True, space_after=4)
add_bullet(doc, "Fev 2026: Expansão do motor de recomendação, integração com "
                "OpenAI e pipeline RAG inicial.")
add_bullet(doc, "Mar 2026: Desenvolvimento do frontend completo, sistema de "
                "design premium, dark/light mode.")
add_bullet(doc, "Abr 2026: Testes automatizados, modernização do código "
                "(remoção de deprecations), deploy em produção. "
                "Entrega intercalar (o presente relatório).")
add_bullet(doc, "Mai 2026: Validação por terceiros, refinamentos finais, "
                "documentação completa. Entrega final e defesa.")

add_para(doc,
    "A Figura 7 apresenta o cronograma Gantt detalhado. À data desta entrega "
    "intercalar, todos os marcos até ao final de abril foram cumpridos dentro "
    "do prazo planeado. As dificuldades mais relevantes encontradas foram:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "Migração do SQLAlchemy v1 para v2 — implicou refactor da "
                "camada de dados, mas produziu código mais limpo e tipado.")
add_bullet(doc, "Modernização para Pydantic v2 — alteração de @validator para "
                "@model_validator e de orm_mode para ConfigDict.")
add_bullet(doc, "Dimensionamento do modelo de dados — iteração sobre o "
                "diagrama ER com base em casos de uso concretos.")
add_bullet(doc, "Configuração CORS em produção — ajuste de origens permitidas "
                "para o domínio Vercel.")

add_para(doc,
    "[Espaço reservado para Figura 7 — Cronograma Gantt. A gerar em "
    "ferramenta externa (TeamGantt, MS Project) e incluir antes da entrega "
    "final.]",
    italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ===========================================================================
# 7. RESULTADOS
# ===========================================================================

add_heading(doc, "Resultados", level=1, number="7")

add_heading(doc, "Resultados dos Testes", level=2, number="7.1")
add_para(doc,
    "A execução da suite de testes automatizados à data desta entrega "
    "intercalar produz os seguintes resultados: 62 testes, 62 aprovados, 0 "
    "falhas, 0 warnings. O tempo total de execução é inferior a 10 segundos, "
    "graças à utilização de SQLite in-memory.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_code_block(doc,
    "$ pytest --tb=short -q\n"
    "............................................................\n"
    "62 passed in 8.43s")

add_para(doc,
    "Os testes de aceitação por terceiros encontram-se em fase de planeamento. "
    "O inquérito, guião de entrevista e resultados serão incluídos em anexo "
    "na entrega final, conforme orientação do template.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Cumprimento de Requisitos", level=2, number="7.2")
add_para(doc,
    "A Tabela 8 sumariza o estado de cumprimento de cada requisito identificado "
    "na Secção 3.1. Os requisitos com estado \"Parcialmente realizado\" "
    "correspondem a funcionalidades com núcleo implementado mas com "
    "refinamentos previstos para a entrega final. Os \"Planeados\" são "
    "requisitos com baixa prioridade que dependem da disponibilidade temporal "
    "até à entrega final.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_caption(doc, "Cumprimento de requisitos.", kind="Tabela 8")
add_table(doc,
    ["ID", "Requisito", "Estado", "Justificação"],
    [
        ["RF01", "Registo e autenticação", "Realizado", "—"],
        ["RF02", "Gestão de perfil", "Realizado", "—"],
        ["RF03", "Registo de treinos", "Realizado", "—"],
        ["RF04", "Registo de refeições", "Realizado", "—"],
        ["RF05", "Tracking hidratação", "Realizado", "—"],
        ["RF06", "Tracking de peso", "Realizado", "—"],
        ["RF07", "Planos de treino e nutrição", "Realizado", "—"],
        ["RF08", "Plano diário adaptativo", "Parcialmente realizado",
         "Adaptação básica implementada; refinamento de triggers pendente."],
        ["RF09", "Chat IA conversacional", "Realizado", "—"],
        ["RF10", "Relatórios e gráficos", "Realizado", "—"],
        ["RF11", "Sessões de meditação", "Realizado", "—"],
        ["RF12", "Categorias personalizáveis", "Realizado", "—"],
        ["RF13", "Pipeline RAG sobre PDFs", "Parcialmente realizado",
         "Extração e consulta funcionais; UI de upload em refinamento."],
        ["RF14", "Notificações push", "Planeado",
         "Dependente de Service Worker; previsto para a entrega final."],
        ["RF15", "Integração wearables", "Não realizado",
         "Adiado para trabalhos futuros — requer tempo adicional de integração."],
        ["RNF01-RNF08", "Não-funcionais principais", "Realizado", "—"],
        ["RNF09", "Internacionalização", "Parcialmente realizado",
         "Mensagens isoladas; sistema i18n completo previsto pós-TFC."],
        ["RNF10", "Privacidade", "Realizado", "—"],
    ],
    col_widths=[1.5, 4.5, 3.0, 6.5])

# ===========================================================================
# 8. CONCLUSÃO
# ===========================================================================

add_heading(doc, "Conclusão", level=1, number="8")

add_heading(doc, "Conclusão", level=2, number="8.1")
add_para(doc,
    "O presente relatório intercalar documenta o estado de desenvolvimento do "
    "LAPHIS à data de abril de 2026, cobrindo os trabalhos desde o início da "
    "UC de Projeto I em setembro de 2025. O problema identificado — a "
    "fragmentação das ferramentas de tracking de saúde e a ausência de "
    "recomendações verdadeiramente personalizadas — foi abordado através da "
    "conceção e implementação de uma plataforma web integrada, que combina "
    "tracking multi-domínio, um motor de recomendação próprio baseado em "
    "heurísticas clínicas reconhecidas, e uma camada opcional de IA "
    "conversacional com pipeline RAG.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "Os principais resultados desta fase incluem: (i) 13 dos 15 requisitos "
    "funcionais encontram-se realizados, os restantes em estado parcial ou "
    "planeado; (ii) 9 dos 10 requisitos não-funcionais cumpridos, um parcial; "
    "(iii) 62 testes automatizados a passar com 0 warnings; (iv) aplicação "
    "em fase de deploy em produção nos serviços Vercel e Render; "
    "(v) aproximadamente 20.300 linhas de código distribuídas entre backend "
    "(~6.900 linhas) e frontend (~13.400 linhas); (vi) integração bem-sucedida "
    "da camada de IA com mecanismo de fallback para modo baseado em regras.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para(doc,
    "As principais conclusões das várias fases do desenvolvimento são: a "
    "escolha do stack tecnológico (FastAPI + React + SQLAlchemy 2.0) revelou-se "
    "acertada, proporcionando alta produtividade e um ciclo de desenvolvimento "
    "rápido; a arquitetura modular facilitou significativamente a expansão "
    "incremental de funcionalidades; a separação clara entre motor de regras "
    "e camada de IA provou-se crítica para garantir o funcionamento do sistema "
    "em ambientes sem chave OpenAI; o investimento em testes automatizados "
    "desde cedo preveniu diversos bugs de regressão durante a modernização do "
    "código.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "Trabalhos Futuros", level=2, number="8.2")
add_para(doc,
    "Até à entrega final de Projeto II, prevista para maio de 2026, serão "
    "realizados os seguintes trabalhos:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "Conclusão do requisito RF08 — refinamento dos triggers de "
                "adaptação automática do plano diário.")
add_bullet(doc, "Conclusão do requisito RF13 — finalização da interface de "
                "upload de PDFs e consulta em linguagem natural.")
add_bullet(doc, "Implementação do requisito RF14 — notificações push via "
                "Service Worker.")
add_bullet(doc, "Aplicação do inquérito de validação a 10-15 utilizadores-teste "
                "e inclusão dos resultados em anexo.")
add_bullet(doc, "Inclusão de screenshots reais dos ecrãs e diagramas UML/ER "
                "exportados de ferramentas externas.")
add_bullet(doc, "Gravação do vídeo demonstrativo e publicação do link no "
                "repositório.")
add_bullet(doc, "Preparação de credenciais de demonstração e documentação de "
                "acesso para os avaliadores.")

add_para(doc,
    "Para além do TFC, identificam-se as seguintes direções de evolução "
    "futura, já fora do âmbito académico:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
add_bullet(doc, "Integração com wearables — Apple Health, Fitbit, Google Fit "
                "via APIs oficiais (RF15).")
add_bullet(doc, "Gamificação — sistema de pontos, badges e streaks para "
                "aumentar retenção e adesão.")
add_bullet(doc, "Progressive Web App — modo offline completo com Service "
                "Worker e sincronização diferida.")
add_bullet(doc, "Machine Learning — substituição progressiva das heurísticas "
                "por modelos treinados nos dados reais dos utilizadores.")
add_bullet(doc, "Aplicação móvel nativa — port para React Native com partilha "
                "do código de lógica de negócio.")
add_bullet(doc, "Marketplace profissional — permitir a nutricionistas e "
                "personal trainers certificados publicarem planos.")
add_bullet(doc, "Scan de código de barras para dados nutricionais automáticos "
                "via API Open Food Facts.")
add_bullet(doc, "Internacionalização completa (i18n) — suporte para inglês e "
                "espanhol no mínimo.")

# ===========================================================================
# BIBLIOGRAFIA
# ===========================================================================

add_heading(doc, "Bibliografia", level=1, number=None)

refs = [
    ("[CiAl24]",
     "B. P. Cipriano e P. Alves, Seven Years Later: Lessons Learned in Automated "
     "Assessment, 5th International Computer Programming Education Conference, "
     "ICPEC 2024, Lisboa, Portugal, jun. 2024."),
    ("[DEISI24]",
     "DEISI, Regulamento de Trabalho Final de Curso, Universidade Lusófona, "
     "set. 2024."),
    ("[FaAP23]",
     "FastAPI Documentation, FastAPI — Modern, fast web framework for building "
     "APIs, 2023. https://fastapi.tiangolo.com/ (website)"),
    ("[HaBe18]",
     "J. A. Harris e F. G. Benedict, A Biometric Study of Human Basal Metabolism, "
     "PNAS, 1918."),
    ("[OpAI24]",
     "OpenAI, OpenAI Platform Documentation, 2024. "
     "https://platform.openai.com/docs (website)"),
    ("[ReAc24]",
     "React Documentation, React — The library for web and native user interfaces, "
     "Meta, 2024. https://react.dev/ (website)"),
    ("[SQLA23]",
     "SQLAlchemy Documentation, SQLAlchemy 2.0 — The Python SQL Toolkit and "
     "ORM, 2023. https://www.sqlalchemy.org/ (website)"),
    ("[TaWe20]",
     "A. Tanenbaum e D. Wetherall, Computer Networks, 6ª Edição, Prentice Hall, "
     "USA, 2020. (livro)"),
    ("[ULHT24]",
     "Universidade Lusófona de Humanidades e Tecnologias, www.ulusofona.pt, "
     "out. 2024. (website)"),
    ("[VeRe24]",
     "Vercel Documentation, Vercel — Develop. Preview. Ship., 2024. "
     "https://vercel.com/docs (website)"),
    ("[WHO24]",
     "World Health Organization, Physical Activity Guidelines, 2024. "
     "https://www.who.int/ (website)"),
]
for key, text in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(2.0)
    p.paragraph_format.first_line_indent = Cm(-2.0)
    run = p.add_run(key)
    set_run(run, size=10.5, bold=True)
    run = p.add_run(f"\t{text}")
    set_run(run, size=10.5)

# ===========================================================================
# ANEXO 1
# ===========================================================================

add_heading(doc, "Anexo 1 — Recomendações para formatação do relatório",
            level=1, number=None)

add_para(doc,
    "Este anexo segue a orientação do template original da Universidade "
    "Lusófona, documentando convenções de formatação aplicadas neste relatório "
    "e servindo de referência para revisões futuras.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc, "Formatação geral", bold=True, space_after=4)
add_bullet(doc, "Fonte do corpo do texto: Calibri 11 pt, espaçamento entre "
                "linhas 1.15.")
add_bullet(doc, "Margens: 2.5 cm (topo, base, direita); 3.0 cm (esquerda) "
                "para encadernação.")
add_bullet(doc, "Cabeçalho: título do trabalho em itálico, alinhado à direita.")
add_bullet(doc, "Rodapé: número de página alinhado à direita.")
add_bullet(doc, "Títulos: Calibri bold, cor azul corporativo (#1F3A68), com "
                "numeração hierárquica (1, 1.1, 1.1.1).")

add_para(doc, "Figuras e tabelas", bold=True, space_after=4)
add_bullet(doc, "Numeradas sequencialmente (Figura 1, Figura 2, ...; Tabela 1, "
                "Tabela 2, ...).")
add_bullet(doc, "Legendas em Calibri 9.5 pt itálico, centradas, imediatamente "
                "abaixo da figura ou acima da tabela.")
add_bullet(doc, "Devem ser sempre referenciadas no texto antes de aparecerem "
                "(\"ver Figura 1\", \"conforme apresentado na Tabela 2\").")

add_para(doc, "Referências bibliográficas", bold=True, space_after=4)
add_bullet(doc, "Formato: chave com 4 caracteres iniciais dos nomes dos "
                "autores + 2 dígitos do ano — ex.: [TaWe20], [OpAI24].")
add_bullet(doc, "Ordenação alfabética pela chave.")
add_bullet(doc, "Citações no texto entre parêntesis retos — ex.: \"segundo "
                "[TaWe20], ...\".")

add_para(doc,
    "Recomenda-se a utilização das funcionalidades \"Referências → Índice\" e "
    "\"Referências → Inserir Legenda\" do MS Word para manter numeração e "
    "paginação automaticamente atualizadas. Após edições substanciais, "
    "selecionar todo o documento e premir F9 para regenerar índices.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ===========================================================================
# SAVE
# ===========================================================================

doc.save(OUTPUT)
print(f"Relatório gerado: {OUTPUT}")

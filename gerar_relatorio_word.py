"""
Gera o relatorio LAPHIS Projeto II em formato Word (.docx)
Formato profissional, pronto para entregar.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

OUTPUT = "docs/LAPHIS_Relatorio_ProjetoII.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = doc.add_paragraph(style=style_map.get(level, "Heading 1"))
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)  # azul
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    else:
        run.font.size = Pt(11)
        run.font.bold = True
    return p

def add_body(doc, text, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_code(doc, text):
    p = doc.add_paragraph(style="No Spacing")
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p.paragraph_format.left_indent = Cm(0.8)
    # light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    return p

def add_table_from_rows(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
        # header bg
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "DBEAFE")
        tcPr.append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def page_break(doc):
    doc.add_page_break()

# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

# Default body font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ---------------------------------------------------------------------------
# CAPA
# ---------------------------------------------------------------------------

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPHIS")
r.font.name = "Calibri"
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Life and Physical Health Intelligent System")
r.font.size = Pt(13)
r.font.italic = True
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Relatório de Projeto II")
r.font.size = Pt(18)
r.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Engenharia Informática — Universidade de Trás-os-Montes e Alto Douro")
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("22 de Abril de 2026")
r.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Grupo: M4ster32")
r.font.size = Pt(11)
r.font.bold = True

page_break(doc)

# ---------------------------------------------------------------------------
# RESUMO
# ---------------------------------------------------------------------------

add_heading(doc, "Resumo", 1)
add_body(doc, (
    "O LAPHIS (Life and Physical Health Intelligent System) é uma plataforma web completa de gestão "
    "de saúde e fitness com recomendações personalizadas por Inteligência Artificial. A aplicação integra "
    "tracking de treinos, nutrição, hidratação, peso corporal, meditação e planeamento diário adaptativo, "
    "com um sistema de chat conversacional com IA."
))
add_body(doc, (
    "O backend é implementado em FastAPI com SQLAlchemy e SQLite/PostgreSQL, enquanto o frontend é "
    "uma SPA React 19 com Vite. O projeto inclui deploy em produção na Vercel (frontend) e Render (backend), "
    "uma suite de 62 testes automatizados e um motor de recomendação de 857 linhas."
))

page_break(doc)

# ---------------------------------------------------------------------------
# 1. INTRODUÇÃO
# ---------------------------------------------------------------------------

add_heading(doc, "1. Introdução", 1)
add_body(doc, (
    "O LAPHIS surgiu da necessidade de centralizar o acompanhamento da saúde e fitness numa única plataforma "
    "inteligente. Enquanto as soluções existentes (MyFitnessPal, Strava, Calm) funcionam em silos, o LAPHIS "
    "integra todos os domínios e utiliza IA para cruzar os dados e gerar recomendações verdadeiramente personalizadas."
))

add_heading(doc, "1.1. Objetivos", 2)
for obj in [
    "Implementar uma plataforma web completa de saúde e fitness",
    "Integrar IA para recomendações personalizadas e chat conversacional",
    "Garantir qualidade com testes automatizados e código sem deprecations",
    "Realizar deploy em produção com CI/CD automático",
    "Criar um sistema de design premium com dark/light mode",
]:
    add_bullet(doc, obj)

add_heading(doc, "1.2. Estrutura do Projeto", 2)
add_code(doc, "laphis/")
add_code(doc, "├── ai-service/          # Backend FastAPI")
add_code(doc, "│   └── src/")
add_code(doc, "│       ├── api/         # 17 routers (auth, logs, plans, chat, ...)")
add_code(doc, "│       ├── core/        # DB, models, schemas, recommender")
add_code(doc, "│       └── utils/       # Email, helpers")
add_code(doc, "└── laphis-frontend/     # Frontend React 19")
add_code(doc, "    └── src/")
add_code(doc, "        ├── pages/       # 17 páginas")
add_code(doc, "        ├── components/  # 12 componentes reutilizáveis")
add_code(doc, "        └── services/    # Service layer (api.js)")

page_break(doc)

# ---------------------------------------------------------------------------
# 2. ARQUITETURA
# ---------------------------------------------------------------------------

add_heading(doc, "2. Arquitetura do Sistema", 1)
add_body(doc, (
    "O LAPHIS segue uma arquitetura cliente-servidor com separação clara entre frontend SPA e backend RESTful API. "
    "A comunicação é feita via HTTP/JSON com autenticação JWT stateless."
))

add_code(doc, "Browser  →  Vercel (React SPA)  →  Render (FastAPI)  →  PostgreSQL")
doc.add_paragraph()

add_heading(doc, "2.1. Backend — FastAPI", 2)
for item in [
    "17 routers registados, cada um com prefixo e tags próprios",
    "Lifespan handler para inicialização da base de dados",
    "CORS configurado para aceitar pedidos do frontend (Vercel)",
    "Documentação automática OpenAPI em /docs",
    "68 endpoints no total",
]:
    add_bullet(doc, item)

add_heading(doc, "2.2. Frontend — React 19 SPA", 2)
for item in [
    "17 páginas com routing protegido por JWT (AppRouter.jsx)",
    "Context API para estado global (AppContext, ThemeContext)",
    "Service layer centralizado em services/api.js",
    "Componentes reutilizáveis: Button, Card, Modal, Form, Toast, Skeleton",
]:
    add_bullet(doc, item)

page_break(doc)

# ---------------------------------------------------------------------------
# 3. TECNOLOGIAS
# ---------------------------------------------------------------------------

add_heading(doc, "3. Tecnologias Utilizadas", 1)

add_heading(doc, "3.1. Backend", 2)
add_table_from_rows(doc,
    ["Tecnologia", "Versão", "Propósito"],
    [
        ["Python", "3.11+", "Linguagem principal"],
        ["FastAPI", "0.104.1", "Framework web async com OpenAPI automático"],
        ["SQLAlchemy", "2.0.23", "ORM com query builder e suporte a migrations"],
        ["Pydantic", "2.5.0", "Validação de dados e serialização"],
        ["Uvicorn", "0.24.0", "Servidor ASGI de produção"],
        ["PyJWT", "2.12.1", "Gestão de tokens JWT stateless"],
        ["bcrypt", "4.1.2", "Hashing seguro de passwords com salt"],
        ["OpenAI", "1.14.0", "API GPT para funcionalidades RAG"],
        ["PyMuPDF", "1.24.0", "Extração de texto PDF para pipeline RAG"],
    ]
)

add_heading(doc, "3.2. Frontend", 2)
add_table_from_rows(doc,
    ["Tecnologia", "Versão", "Propósito"],
    [
        ["React", "19.2.0", "Biblioteca UI com concurrent rendering"],
        ["Vite", "8.0 (beta)", "Build tool com HMR instantâneo via ESBuild"],
        ["React Router", "7.13.1", "Routing declarativo com layouts aninhados"],
        ["Lucide React", "0.577.0", "300+ ícones SVG consistentes"],
        ["Recharts", "3.8.0", "Gráficos React com SVG responsivo"],
        ["jsPDF", "4.2.0", "Geração de relatórios PDF no cliente"],
    ]
)

add_heading(doc, "3.3. Infraestrutura", 2)
add_table_from_rows(doc,
    ["Ferramenta", "Propósito"],
    [
        ["Vercel", "Hosting frontend com CI/CD automático"],
        ["Render", "Hosting backend com PostgreSQL gerido"],
        ["GitHub", "Controlo de versões e colaboração"],
        ["pytest", "Suite de testes com fixtures e parametrização"],
        ["ESLint", "Linting do código JavaScript/React"],
    ]
)

page_break(doc)

# ---------------------------------------------------------------------------
# 4. MODELO DE DADOS
# ---------------------------------------------------------------------------

add_heading(doc, "4. Modelo de Dados", 1)
add_body(doc, (
    "O modelo de dados consiste em 16 entidades SQLAlchemy, todas referenciando o utilizador (User) "
    "como entidade central. Todas as datas são timezone-aware e os campos semi-estruturados usam JSON fields."
))

add_heading(doc, "4.1. Entidades Principais", 2)
add_table_from_rows(doc,
    ["Entidade", "Campos Relevantes", "Propósito"],
    [
        ["User", "email, hashed_pw, is_verified, verify_code, reset_token", "Autenticação e identidade"],
        ["Profile", "age, weight, height, goal, activity, diet_type", "Dados biométricos do utilizador"],
        ["Plan", "title, type, content, is_active, feedback", "Planos de treino/nutrição gerados pela IA"],
        ["ChatSession / ChatMessage", "title, role, content", "Conversas com o assistente IA"],
        ["WorkoutLog", "exercise, duration, calories, category", "Registos de treino"],
        ["MealLog", "food_name, calories, protein, carbs, fat, meal_type", "Registos de refeição"],
        ["WaterLog", "amount_ml, logged_at", "Tracking de hidratação"],
        ["WeightEntry", "weight_kg, logged_at", "Tracking de peso corporal"],
        ["ZenSession", "type, duration_min, mood_before, mood_after, notes", "Sessões de meditação"],
        ["Category", "name, type, icon, color", "Categorias personalizáveis"],
        ["WeeklySummary", "week_start, avg_calories, total_workouts, recommendations", "Resumos semanais"],
        ["DailyPlan", "date, plan_data (JSON), status", "Planeamento diário adaptativo"],
        ["ProgressSnapshot", "date, metrics (JSON)", "Snapshots de progresso"],
        ["AdaptationLog", "trigger, old_params, new_params, reason", "Log de adaptações da IA"],
        ["PlanFeedback", "plan_id, rating, comment", "Feedback dos utilizadores sobre planos"],
    ]
)

page_break(doc)

# ---------------------------------------------------------------------------
# 5. BACKEND — IMPLEMENTAÇÃO
# ---------------------------------------------------------------------------

add_heading(doc, "5. Implementação do Backend", 1)

add_heading(doc, "5.1. Autenticação", 2)
add_body(doc, "Fluxo completo e seguro:")
for step in [
    "Registo: email + password → hash bcrypt → código de verificação (6 dígitos) por SMTP",
    "Verificação de email: utilizador insere código → conta ativada (is_verified = True)",
    "Login: email + password → verificação bcrypt → JWT com expiração de 7 dias",
    "Reset de password: email → token único → novo hash guardado",
    "Middleware: cada endpoint protegido valida o JWT do header Authorization: Bearer <token>",
]:
    add_bullet(doc, step)

add_heading(doc, "5.2. Módulos Implementados", 2)
add_table_from_rows(doc,
    ["Módulo", "Endpoints", "Funcionalidade"],
    [
        ["auth.py", "POST /register, /login, /verify, /reset", "Autenticação completa"],
        ["profile.py", "GET/PUT /profile", "Gestão de dados biométricos"],
        ["logs.py", "CRUD /logs", "Treinos e refeições (schema unificado)"],
        ["water.py", "POST/GET /water", "Tracking de hidratação"],
        ["weight.py", "POST/GET /weight", "Tracking de peso e BMI"],
        ["plans.py", "CRUD /plans + /generate + /feedback", "Planos gerados pela IA"],
        ["chat.py", "CRUD /chat/sessions + /messages", "Chat conversacional com IA"],
        ["zen.py", "POST/GET /zen + /stats", "Sessões de meditação"],
        ["reports.py", "GET /reports/weekly + /summary", "Resumos e exportação"],
        ["categories.py", "CRUD /categories", "Categorias personalizáveis"],
        ["ingest.py", "POST /ingest", "Upload de PDFs para RAG"],
    ]
)

add_heading(doc, "5.3. Motor de Recomendação (recommender.py)", 2)
add_body(doc, (
    "O recommender.py tem 857 linhas e implementa lógica rule-based com heurísticas nutricionais e de exercício, "
    "complementada pela API OpenAI para funcionalidades RAG."
))
for item in [
    "BMI: peso / altura²",
    "TDEE: equação de Harris-Benedict com multiplicador de atividade",
    "Macronutrientes: distribuição personalizada por objetivo (perda peso / ganho muscular / manutenção)",
    "Hidratação: recomendação de 30-35ml/kg",
    "Planeamento diário: treino + refeições com macros calculados, adaptável",
    "Pipeline RAG: PyMuPDF → indexação → OpenAI GPT contextualizado",
]:
    add_bullet(doc, item)

page_break(doc)

# ---------------------------------------------------------------------------
# 6. FRONTEND
# ---------------------------------------------------------------------------

add_heading(doc, "6. Implementação do Frontend", 1)

add_heading(doc, "6.1. Routing e Páginas", 2)
add_table_from_rows(doc,
    ["Rota", "Página", "Acesso"],
    [
        ["/", "Home (landing page)", "Público"],
        ["/login, /register", "Login / Registo", "Público"],
        ["/verify-email, /forgot-password, /reset-password", "Auth flows", "Público"],
        ["/app", "Dashboard", "Protegido (JWT)"],
        ["/app/logs", "Logs (treinos e refeições)", "Protegido"],
        ["/app/plans, /app/plans/:id", "Planos + Detalhe", "Protegido"],
        ["/app/chat", "Chat com IA", "Protegido"],
        ["/app/zen", "Meditação guiada", "Protegido"],
        ["/app/reports", "Relatórios e gráficos", "Protegido"],
        ["/app/profile, /app/settings", "Perfil e Definições", "Protegido"],
    ]
)

add_heading(doc, "6.2. Componentes Reutilizáveis", 2)
add_table_from_rows(doc,
    ["Componente", "Propósito"],
    [
        ["Button", "Variantes primary, secondary, ghost, danger — com loading state e ícone opcional"],
        ["Card", "Container visual com elevation levels e hover effects"],
        ["Modal", "Diálogo sobreposto com backdrop blur e animações"],
        ["Form", "Inputs estilizados com validação e estados de erro"],
        ["Toast", "Notificações temporárias (success, error, warning, info)"],
        ["Skeleton", "Loading placeholders com animação pulse"],
        ["Icon", "Wrapper de ícones Lucide com tamanhos padronizados"],
    ]
)

add_heading(doc, "6.3. Gestão de Estado", 2)
for item in [
    "AppContext: autenticação (token JWT), perfil do utilizador, loading states, error handling",
    "ThemeContext: dark/light mode com persistência em localStorage e CSS Custom Properties",
    "Service layer (api.js): comunicação HTTP centralizada com header Authorization automático",
]:
    add_bullet(doc, item)

page_break(doc)

# ---------------------------------------------------------------------------
# 7. SISTEMA DE DESIGN
# ---------------------------------------------------------------------------

add_heading(doc, "7. Sistema de Design", 1)
add_body(doc, (
    "Filosofia 'Premium Monochromatic', inspirada em Linear, Notion e Stripe. "
    "Implementado via CSS Custom Properties com suporte completo a dark/light mode."
))

add_heading(doc, "7.1. Paleta de Cores", 2)
add_table_from_rows(doc,
    ["Variável", "Claro", "Escuro"],
    [
        ["Background primário", "#FFFFFF", "#0A0A0A"],
        ["Background secundário", "#F8F9FA", "#141414"],
        ["Texto primário", "#111111", "#F9FAFB"],
        ["Texto secundário", "#6B7280", "#9CA3AF"],
        ["Acento (CTA)", "#2563EB", "#3B82F6"],
        ["Bordas", "rgba(0,0,0,0.08)", "rgba(255,255,255,0.08)"],
    ]
)

add_heading(doc, "7.2. Tipografia e Espaçamento", 2)
add_table_from_rows(doc,
    ["Escala", "Tamanho", "Weight"],
    [
        ["Display", "36px", "700"],
        ["Title", "24px", "600"],
        ["Subtitle", "18px", "600"],
        ["Body", "15px", "400"],
        ["Caption", "13px", "500"],
        ["Micro", "12px", "500"],
    ]
)
add_body(doc, "Espaçamento em múltiplos de 4px (4px → 8px → 16px → 24px → 32px → 48px).")
add_body(doc, "Responsividade mobile-first: < 640px (single-column), 640-1024px (grid adaptativo), > 1024px (sidebar).")

page_break(doc)

# ---------------------------------------------------------------------------
# 8. TESTES
# ---------------------------------------------------------------------------

add_heading(doc, "8. Testes e Qualidade de Código", 1)
add_body(doc, "Suite de 62 testes automatizados usando pytest e httpx, com base de dados SQLite in-memory para isolamento total.")

add_table_from_rows(doc,
    ["Ficheiro", "Testes", "Cobertura"],
    [
        ["test_auth.py", "~10", "Registo, login, verificação email, reset, tokens inválidos"],
        ["test_profile.py", "~8", "Criar, atualizar, obter perfil, validação de campos"],
        ["test_logs.py", "~12", "Treinos, refeições, listagem, edição, eliminação, filtros"],
        ["test_water.py", "~8", "Registar água, total diário, histórico, meta"],
        ["test_weight.py", "~8", "Registar peso, histórico, tendência, BMI"],
        ["test_plans.py", "~10", "Gerar, guardar, listar, atualizar, eliminar, feedback"],
        ["test_zen.py", "~6", "Criar sessão, listar, estatísticas, tipos"],
    ]
)

add_code(doc, "$ pytest --tb=short -q")
add_code(doc, "62 passed, 0 warnings")
doc.add_paragraph()

add_heading(doc, "8.1. Modernização de Código", 2)
add_table_from_rows(doc,
    ["Padrão Antigo", "Padrão Moderno"],
    [
        ["@app.on_event('startup')", "lifespan context manager (FastAPI 0.100+)"],
        ["Pydantic class Config com orm_mode", "model_config = ConfigDict(from_attributes=True)"],
        ["datetime.utcnow()", "datetime.now(datetime.timezone.utc)"],
        ["@validator", "@model_validator(mode='before') — Pydantic v2"],
    ]
)

page_break(doc)

# ---------------------------------------------------------------------------
# 9. DEPLOY
# ---------------------------------------------------------------------------

add_heading(doc, "9. Deploy e Infraestrutura", 1)

add_heading(doc, "9.1. Arquitetura de Deploy", 2)
add_code(doc, "Browser  -->  Vercel (React SPA)  -->  Render (FastAPI + PostgreSQL)")
doc.add_paragraph()

add_heading(doc, "9.2. Frontend — Vercel", 2)
for item in [
    "Deploy automático: push para main → build → deploy em ~30 segundos",
    "CDN global: assets servidos a partir do edge node mais próximo",
    "HTTPS automático com certificado SSL provisionado",
    "Preview deployments: cada PR gera URL de preview único",
]:
    add_bullet(doc, item)

add_heading(doc, "9.3. Backend — Render", 2)
for item in [
    "Web Service com Uvicorn a servir FastAPI",
    "PostgreSQL gerido pelo Render",
    "Health checks no endpoint /health para auto-restart",
    "Logs em tempo real via dashboard Render",
]:
    add_bullet(doc, item)

add_heading(doc, "9.4. Variáveis de Ambiente", 2)
add_table_from_rows(doc,
    ["Variável", "Contexto", "Propósito"],
    [
        ["DATABASE_URL", "Backend (prod)", "Connection string PostgreSQL"],
        ["JWT_SECRET", "Backend", "Chave de assinatura dos tokens"],
        ["SMTP_SERVER / SMTP_EMAIL / SMTP_PASSWORD", "Backend", "Email para verificação de conta"],
        ["OPENAI_API_KEY", "Backend", "Chave API para funcionalidades RAG"],
        ["VITE_API_URL", "Frontend", "URL do backend para API calls"],
    ]
)

page_break(doc)

# ---------------------------------------------------------------------------
# 10. REQUISITOS
# ---------------------------------------------------------------------------

add_heading(doc, "10. Análise de Requisitos", 1)

add_heading(doc, "10.1. Requisitos Funcionais — 12/12 Implementados (100%)", 2)
add_table_from_rows(doc,
    ["ID", "Requisito", "Estado"],
    [
        ["RF01", "Registo e autenticação de utilizadores", "Completo"],
        ["RF02", "Gestão de perfil com dados biométricos", "Completo"],
        ["RF03", "Registo de treinos e refeições", "Completo"],
        ["RF04", "Tracking de hidratação", "Completo"],
        ["RF05", "Tracking de peso corporal", "Completo"],
        ["RF06", "Recomendações personalizadas com IA", "Completo"],
        ["RF07", "Chat conversacional com assistente IA", "Completo"],
        ["RF08", "Planos de treino e nutrição", "Completo"],
        ["RF09", "Relatórios e visualização de progresso", "Completo"],
        ["RF10", "Sessões de meditação/mindfulness", "Completo"],
        ["RF11", "Categorias personalizáveis", "Completo"],
        ["RF12", "Planeamento diário adaptativo", "Completo"],
    ]
)

add_heading(doc, "10.2. Requisitos Não-Funcionais — 8/8 Implementados", 2)
add_table_from_rows(doc,
    ["ID", "Requisito", "Estado"],
    [
        ["RNF01", "Performance — resposta < 2s", "FastAPI async + SQLAlchemy optimizado"],
        ["RNF02", "Segurança — passwords encriptadas, JWT", "bcrypt + PyJWT + HTTPS em produção"],
        ["RNF03", "Responsividade — mobile e desktop", "CSS mobile-first + breakpoints"],
        ["RNF04", "Disponibilidade — 99%+ uptime", "Vercel + Render com health checks"],
        ["RNF05", "Testabilidade — suite de testes", "62 testes, 0 warnings"],
        ["RNF06", "Manutenibilidade — código limpo", "Separação em camadas, sem deprecations"],
        ["RNF07", "Acessibilidade — dark mode, contraste", "WCAG AA ratios, dark/light toggle"],
        ["RNF08", "Escalabilidade — arquitetura modular", "Stateless API, DB substituível"],
    ]
)

page_break(doc)

# ---------------------------------------------------------------------------
# 11. MÉTRICAS E ESTADO ATUAL
# ---------------------------------------------------------------------------

add_heading(doc, "11. Estado Atual e Métricas", 1)

add_table_from_rows(doc,
    ["Métrica", "Valor"],
    [
        ["Linhas de código — Backend", "~6.900 linhas Python"],
        ["Linhas de código — Frontend", "~13.400 linhas JSX/JS/CSS"],
        ["Total de linhas", "~20.300 linhas"],
        ["Modelos SQLAlchemy", "16 entidades"],
        ["Schemas Pydantic", "30+ schemas"],
        ["Endpoints API", "68 endpoints"],
        ["Routers FastAPI", "17 routers"],
        ["Páginas React", "17 páginas"],
        ["Componentes React", "12 componentes reutilizáveis"],
        ["Testes automatizados", "62 testes (100% pass, 0 warnings)"],
        ["Repositório", "github.com/M4ster32/Laphis.git (branch: main)"],
    ]
)

add_heading(doc, "11.1. Funcionalidades Implementadas", 2)
for feat in [
    "Autenticação completa (registo, login, verificação email, reset password)",
    "Perfil biométrico com cálculos automáticos (BMI, TDEE, macros)",
    "Tracking de treinos e refeições com categorias personalizáveis",
    "Tracking de hidratação com meta diária baseada no peso",
    "Tracking de peso com tendências e cálculo de BMI",
    "Chat conversacional com IA personalizada",
    "Planos de treino e nutrição gerados por IA",
    "Planeamento diário adaptativo",
    "Sessões de meditação/mindfulness com registo de humor",
    "Relatórios semanais e exportação PDF (jsPDF)",
    "Sistema de design premium com dark/light mode completo",
    "Deploy em produção: Vercel (frontend) + Render (backend + PostgreSQL)",
    "Suite de 62 testes automatizados com 0 warnings",
]:
    add_bullet(doc, feat)

page_break(doc)

# ---------------------------------------------------------------------------
# 12. TRABALHO FUTURO
# ---------------------------------------------------------------------------

add_heading(doc, "12. Trabalho Futuro", 1)

add_heading(doc, "12.1. Curto Prazo", 2)
for item in [
    "Notificações push — lembrar utilizador de registar refeições, água e treinos",
    "Gamificação — sistema de pontos, badges e streaks para motivar a consistência",
    "Integração com wearables — Fitbit, Apple Health, Google Fit via APIs",
    "Offline mode — Service Worker para funcionar sem internet (PWA completo)",
]:
    add_bullet(doc, item)

add_heading(doc, "12.2. Médio e Longo Prazo", 2)
for item in [
    "Machine Learning — substituir heurísticas por modelos treinados nos dados dos utilizadores",
    "Social features — partilha de progresso, desafios entre utilizadores",
    "Nutrição avançada — scan de código de barras para dados nutricionais automáticos",
    "Aplicação móvel nativa — React Native para iOS e Android",
    "Marketplace de planos — planos criados por profissionais de saúde certificados",
]:
    add_bullet(doc, item)

page_break(doc)

# ---------------------------------------------------------------------------
# 13. CONCLUSÃO
# ---------------------------------------------------------------------------

add_heading(doc, "13. Conclusão", 1)
add_body(doc, (
    "O projeto LAPHIS atingiu com sucesso todos os objetivos propostos para a fase de implementação (Projeto II). "
    "Os 12 requisitos funcionais e 8 requisitos não-funcionais foram integralmente implementados, resultando "
    "numa aplicação completa e funcional em produção."
))
add_body(doc, (
    "A aplicação encontra-se acessível a utilizadores reais, com frontend deployado na Vercel e backend "
    "na Render com PostgreSQL gerido. A qualidade do código é assegurada por 62 testes automatizados com "
    "0 warnings e pela modernização completa do código para eliminar padrões deprecated."
))
add_body(doc, (
    "As principais conquistas técnicas incluem um motor de recomendação de 857 linhas com lógica nutricional "
    "e de exercício personalizada, um sistema de design premium com dark mode completo, uma arquitetura "
    "limpa e modular com ~20.300 linhas de código, e deploy automatizado com CI/CD integrado."
))

page_break(doc)

# ---------------------------------------------------------------------------
# REFERÊNCIAS
# ---------------------------------------------------------------------------

add_heading(doc, "Referências", 1)
refs = [
    "FastAPI Documentation. (2024). FastAPI — Modern, fast web framework for building APIs. https://fastapi.tiangolo.com/",
    "React Documentation. (2024). React — The library for web and native user interfaces. https://react.dev/",
    "SQLAlchemy Documentation. (2024). SQLAlchemy — The Python SQL Toolkit and ORM. https://www.sqlalchemy.org/",
    "Pydantic Documentation. (2024). Pydantic — Data validation using Python type hints. https://docs.pydantic.dev/",
    "Vite Documentation. (2024). Vite — Next Generation Frontend Tooling. https://vite.dev/",
    "Harris, J.A. & Benedict, F.G. (1918). A Biometric Study of Human Basal Metabolism. PNAS.",
    "World Health Organization. (2024). Physical Activity Guidelines. https://www.who.int/",
    "Vercel Documentation. (2024). Vercel — Develop. Preview. Ship. https://vercel.com/docs",
    "Render Documentation. (2024). Render — Cloud Application Hosting. https://docs.render.com/",
    "OpenAI API Documentation. (2024). OpenAI Platform. https://platform.openai.com/docs",
    "Recharts Documentation. (2024). Recharts — A composable charting library. https://recharts.org/",
    "pytest Documentation. (2024). pytest — helps you write better programs. https://docs.pytest.org/",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}] {ref}")
    run.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(3)

# ---------------------------------------------------------------------------
# SAVE
# ---------------------------------------------------------------------------

doc.save(OUTPUT)
print(f"Word gerado: {OUTPUT}")
